"""Regression tests for DOCX proposal/LOI parsing quality."""

from __future__ import annotations

from io import BytesIO

from docx import Document

import main
from models import Scenario
from scenario_extract import _apply_safe_defaults, _regex_prefill, _should_ocr_docx_images, extract_text_from_docx


def test_extract_text_from_docx_includes_table_rows() -> None:
    doc = Document()
    doc.add_paragraph("Phase-In: Tenant shall pay based on the following RSF.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Months 1 - 12"
    table.rows[0].cells[1].text = "3,300 RSF"
    table.rows[1].cells[0].text = "Months 13 - 24"
    table.rows[1].cells[1].text = "4,600 RSF"
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    text = extract_text_from_docx(buf)
    assert "Months 1 - 12 | 3,300 RSF" in text
    assert "Months 13 - 24 | 4,600 RSF" in text


def test_regex_prefill_parses_docx_proposal_month_phrase_schedule() -> None:
    text = (
        "Premises: Spec Suite 1100 consisting of 5,900 rentable square feet (RSF) on the first (1st) floor of Building 1.\n"
        "Commencement Date: Earlier of occupancy or June 1, 2026.\n"
        "Lease Term: Forty-eight (48) months.\n"
        "Abatement: The first four (4) months of Gross Rent following the Commencement Date shall be abated.\n"
        "Base Rent Schedule: $38.00 per RSF NNN for months 1-12 with 3% annual escalations starting in month 13.\n"
        "Tenant Allowance: Landlord shall provide Tenant with an allowance of up to $5.00 per RSF.\n"
        "Operating Expenses for 2026 are currently estimated to be $21.97 per RSF.\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("rsf") == 5900.0
    assert prefill.get("commencement") == "2026-06-01"
    assert prefill.get("term_months") == 48
    assert prefill.get("free_rent_months") == 4
    assert prefill.get("ti_allowance_psf") == 5.0
    assert prefill.get("base_opex_psf_yr") == 21.97
    assert prefill.get("_rent_steps_source") == "month_phrase_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 38.0}
    assert steps[-1]["end"] == 47


def test_extract_lease_hints_parses_loi_term_suite_and_opex_from_docx_style_text() -> None:
    text = (
        "Building: Eastbound - 3232 E. Cesar Chavez St. Austin, Texas 78702 - Building 2.\n"
        "Premises: Option 1 - 3,226 RSF located in Suite 2-380.\n"
        "Lease Commencement: estimated to be May 1, 2026.\n"
        "Lease Term: Landlord is proposing a 52 month term.\n"
        "Base Rent: Landlord is proposing initial base rent of $42.00 /rsf, net of operating, with 3% annual increases.\n"
        "Operating Expenses: Estimated operating expenses for 2026 are $14.26. 2025 Operating expenses were $14.35.\n"
        "Tenant Improvements: Tenant shall be provided an allowance of $1.50 RSF.\n"
    )
    hints = main._extract_lease_hints(text, "proposal.docx", "test-rid")

    assert hints["building_name"] == "Eastbound - 3232 E. Cesar Chavez St. Austin"
    assert hints["suite"] == "2-380"
    assert hints["rsf"] == 3226.0
    assert str(hints["commencement_date"]) == "2026-05-01"
    assert hints["term_months"] == 52
    assert str(hints["expiration_date"]) == "2030-08-31"
    assert hints["opex_psf_year_1"] == 14.26
    assert hints["opex_source_year"] == 2026
    assert hints["ti_allowance_psf"] == 1.5


def test_extract_lease_hints_handles_reversed_suite_tokens_and_proposal_header_noise() -> None:
    text = (
        "OAKPOINTE\n"
        "9111 Jollyville Road, Austin, TX 78759\n"
        "LEASE PROPOSAL FOR\n"
        "Law Office of Don E. Walden\n"
        "PREMISES | 245 Suite (See attached floor plan)\n"
        "SQUARE FOOTAGE | 1,145 RSF\n"
        "COMMENCEMENT | May 1, 2026\n"
        "TERM | 36 Months | 60 Months | 60 Months\n"
        "BASE RENT | $17.00/RSF | $16.50/RSF | $16.50/RSF\n"
        "ABATEMENT | 1 Month outside the lease term. | 3 Months outside the lease term.\n"
    )
    hints = main._extract_lease_hints(text, "oakpointe-proposal.docx", "test-rid")

    assert hints["suite"] == "245"
    assert hints["address"] == "9111 Jollyville Road, Austin, TX 78759"
    assert "LEASE PROPOSAL FOR" not in str(hints["address"])
    assert "LEASE PROPOSAL FOR" not in str(hints["building_name"])
    assert hints["term_months"] == 60
    schedule = hints.get("rent_schedule")
    assert isinstance(schedule, list)
    assert schedule
    assert schedule[0]["rent_psf_annual"] == 17.0


def test_regex_prefill_prefers_explicit_turnkey_zero_allowance_over_test_fit_allowance() -> None:
    text = (
        'In addition to the "turn-key" delivery, Landlord will provide an allowance equal to $.00 per RSF '
        "that Tenant can use towards moving costs, FF&E, security, and data & cabling.\n"
        "Test-Fit: Landlord will provide a test-fit allowance equal to $0.15 per RSF of the Premises.\n"
    )

    prefill = _regex_prefill(text)

    assert prefill.get("ti_allowance_psf") == 0.0


def test_extract_lease_hints_avoids_false_full_service_from_grossed_up_opex_clause() -> None:
    text = (
        "Re: Office Lease Proposal for Sprinklr (Tenant).\n"
        "Building: 1300 East 5th, located at 1300 E. 5th Street, Austin, TX 78702.\n"
        "Commencement Date: The Commencement Date will upon Substantial Completion of the Premises estimated to be no later than December 1st, 2026.\n"
        "Lease Term: Ninety-one (91) months from the Commencement Date.\n"
        "Base Annual Net Rental Rate: Months 1-12: Annual Base Rent: $43.50/RSF.\n"
        'In addition to the "turn-key" delivery, Landlord will provide an allowance equal to $10.00 per RSF that Tenant can use towards moving costs, FF&E, security, and data & cabling.\n'
        "Operating Expenses and Real Estate Taxes: In addition to the Base Rental Rate, Tenant will be responsible for its proportionate share of Building Operating Expenses and Real Estate Taxes during the term of the lease. 2026 estimated Operating Expenses and Real Estate Taxes are $20.06 per RSF. All operating expenses shall be grossed up to reflect 100% occupancy.\n"
        "Parking: Tenant will have access to a parking ratio of 2.7 spaces per 1,000 RSF. There will be a charge of $185 per month per space for unreserved spaces; reserved spaces are $250 per month per space.\n"
    )

    hints = main._extract_lease_hints(text, "sprinklr-7-year.docx", "test-rid")

    assert hints["building_name"] == "1300 East 5th"
    assert hints["address"] == "1300 E. 5th Street, Austin, TX 78702"
    assert hints["lease_type"] == "NNN"
    assert "full_service_rate_psf_yr" not in hints
    assert hints["opex_psf_year_1"] == 20.06
    assert hints["opex_source_year"] == 2026
    assert hints["ti_allowance_psf"] == 10.0
    assert hints["parking_rate_monthly"] == 185.0


def test_apply_safe_defaults_prefers_document_prefill_for_explicit_ti_and_parking() -> None:
    raw = {
        "scenario": {
            "name": "Sprinklr 1300 E 5th",
            "rsf": 12153,
            "commencement": "2026-12-01",
            "expiration": "2032-05-31",
            "rent_steps": [{"start": 0, "end": 65, "rate_psf_yr": 43.5}],
            "ti_allowance_psf": 5.0,
            "parking_spaces": 7,
            "parking_cost_monthly_per_space": 0.0,
        },
        "confidence": {"ti_allowance_psf": 0.8, "parking_spaces": 0.8},
        "warnings": [],
    }
    prefill = {
        "term_months": 66,
        "ti_allowance_psf": 0.0,
        "parking_spaces": 33,
        "parking_cost_monthly_per_space": 185.0,
    }

    scenario, _confidence, warnings = _apply_safe_defaults(raw, prefill=prefill)

    assert scenario["ti_allowance_psf"] == 0.0
    assert scenario["parking_spaces"] == 33
    assert scenario["parking_cost_monthly_per_space"] == 185.0
    assert any("explicit allowance stated in the document" in warning for warning in warnings)
    assert "Parking count was reset to the document-derived entitlement." in warnings


def test_apply_safe_defaults_normalizes_none_numeric_fields() -> None:
    raw = {
        "scenario": {
            "name": "Example",
            "rsf": 1000,
            "commencement": "2026-01-01",
            "expiration": "2026-12-31",
            "rent_steps": [{"start": 0, "end": 11, "rate_psf_yr": 25.0}],
            "free_rent_months": None,
            "ti_allowance_psf": None,
            "opex_mode": "nnn",
            "base_opex_psf_yr": None,
            "base_year_opex_psf_yr": None,
            "opex_growth": None,
            "discount_rate_annual": None,
        }
    }
    scenario_dict, confidence, warnings = _apply_safe_defaults(raw, prefill={})
    scenario = Scenario.model_validate(scenario_dict)
    assert scenario.free_rent_months == 0
    assert scenario.ti_allowance_psf == 0.0
    assert scenario.base_opex_psf_yr == 10.0
    assert scenario.base_year_opex_psf_yr == 10.0
    assert scenario.opex_growth == 0.03
    assert scenario.discount_rate_annual == 0.08
    assert isinstance(confidence, dict)
    assert isinstance(warnings, list)


def test_regex_prefill_prefers_basic_rent_over_opex_line_in_proposal() -> None:
    text = (
        "Building: Vista Ridge\n"
        "Contraction Premises: A demised portion of Suite 450, approximately 5,944 rentable square feet.\n"
        "Term: Sixty-three (63) Months\n"
        "Commencement Date: December 1, 2026\n"
        "Basic Rent: $27.00 per rentable square foot with 3% annual increases beginning in month 13\n"
        "Rental Abatement: The initial three (3) months' base rent shall be abated\n"
        "Operating Expenses: 2026 Operating Expenses are estimated to be $14.50/sf.\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("rate_psf_yr") == 27.0
    assert prefill.get("base_opex_psf_yr") == 14.5
    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 27.0}
    assert steps[-1]["end"] == 62


def test_regex_prefill_tarrytown_name_ignores_legal_entity_suite_noise() -> None:
    text = (
        "RE: Lease Proposal - Tarrytown Expocare at Research Park Building 3\n"
        "Premises: Research Park Building 3, 12515 Research Park Loop, Austin, TX 78759\n"
        "Landlord Legal Entity: SIR Properties Trust, with an address of 255 Washington Street, Suite 300, Newton, Massachusetts 02458\n"
        "Lease Commencement Date: January 1, 2028\n"
        "Term: Ten (10) years and four (4) months from the Lease Commencement Date.\n"
    )

    prefill = _regex_prefill(text)

    assert prefill.get("term_months") == 124
    assert prefill.get("name") == "Research Park Building 3"


def test_regex_prefill_parses_label_value_annual_base_rent_escalation() -> None:
    text = (
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Lease Commencement: May 1, 2027\n"
        "Lease Expiration: April 30, 2032\n"
        "Base Rent | $42.00/SF\n"
        "Annual Base Rent Escalation | 3.00%\n"
        "Base Operating Expenses | $26.80/SF\n"
        "Annual Opex Escalation | 3.00%\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 42.0}
    assert steps[1]["rate_psf_yr"] == 43.26
    assert steps[-1]["end"] == 58


def test_regex_prefill_parses_annual_rental_increases_with_percent_after_phrase() -> None:
    text = (
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Commencement Date: May 1, 2027.\n"
        "Lease Term: Sixty (60) months from the Commencement Date.\n"
        "Lease Rate: First year Base Rental Rate: $42.00 NNN\n"
        "Annual Rental Increases: The Base Rental Rate shall increase 3% per year.\n"
        "Base Operating Expenses: $26.80/SF.\n"
        "Parking: Per existing lease.\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 42.0}
    assert steps[1]["rate_psf_yr"] == 43.26
    assert steps[-1]["end"] == 59


def test_regex_prefill_parses_word_percent_annual_escalation() -> None:
    text = (
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Commencement Date: May 1, 2027.\n"
        "Lease Term: Sixty (60) months from the Commencement Date.\n"
        "Base Rent: $42.00 / RSF / YR NNN with four percent annual escalations beginning in month 13.\n"
        "Base Operating Expenses: $26.80/SF.\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 42.0}
    assert steps[1] == {"start": 12, "end": 23, "rate_psf_yr": 43.68}
    assert steps[-1]["end"] == 59


def test_regex_prefill_parses_lease_rate_label_without_base_rent_phrase() -> None:
    text = (
        "Premises: Portion of Suite 500 totaling approximately 4,165 rentable square feet.\n"
        "Commencement Date: October 1, 2026.\n"
        "Lease Term: Seventy-eight (78) months.\n"
        "Lease Rate: $44.50 per square foot per year NNN with three percent (2.75%) annual escalations beginning month 19.\n"
        "Operating Expenses: Estimated to be $23.08 per square foot for 2026.\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("rate_psf_yr") == 44.5
    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 17, "rate_psf_yr": 44.5}
    assert steps[1] == {"start": 18, "end": 29, "rate_psf_yr": 45.72}
    assert steps[-1]["rate_psf_yr"] == 50.96
    assert steps[-1]["end"] == 77


def test_extract_lease_hints_builds_escalated_schedule_from_annual_rental_increases_phrase() -> None:
    text = (
        "Building: Lamar Central\n"
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Commencement Date: May 1, 2027.\n"
        "Lease Term: Sixty (60) months from the Commencement Date.\n"
        "Lease Rate: First year Base Rental Rate: $42.00 NNN\n"
        "Annual Rental Increases: The Base Rental Rate shall increase 3% per year.\n"
        "Base Operating Expenses: $26.80/SF.\n"
    )
    hints = main._extract_lease_hints(text, "proposal.docx", "test-rid")
    schedule = hints.get("rent_schedule")
    assert isinstance(schedule, list)


def test_regex_prefill_counterproposal_keeps_opex_ti_and_parking_entitlement_separate() -> None:
    text = (
        "Premises Suite 1370 for approximately 4,064 rentable square feet.\n"
        "Commencement Date September 1, 2026\n"
        "Lease Term Thirty-eight (38) months\n"
        "Rental Abatement Tenant shall have four (4) months of rental abatement. Tenant shall be responsible for the operating expenses during this abatement period.\n"
        "Lease Rate $45.00 per square foot per year NNN with 3% annual escalations.\n"
        "Operating Expenses Tenant shall pay its pro rata share of building expenses, which are estimated to be $24.44 per square foot for 2025 (CAM $12.57, Taxes $11.53, Insurance $0.34).\n"
        "Initial Fit Plan Landlord shall provide an initial fit plan allowance of $0.15/RSF.\n"
        "Tenant Improvements Landlord will pay a tenant allowance of $5.00 per square foot.\n"
        "Parking Tenant shall be entitled to the non-exclusive use of 1 parking space per 600 rentable square feet. The current rate for unreserved parking is $250.00 per space plus tax.\n"
    )

    prefill = _regex_prefill(text)

    assert prefill.get("rsf") == 4064.0
    assert prefill.get("term_months") == 38
    assert prefill.get("rate_psf_yr") == 45.0
    assert prefill.get("base_opex_psf_yr") == 24.44
    assert prefill.get("ti_allowance_psf") == 5.0
    assert prefill.get("parking_ratio_per_1000_rsf") == 1.6667
    assert prefill.get("parking_spaces") == 7


def test_should_ocr_docx_images_skips_clean_counterproposal_text() -> None:
    text = (
        "Premises Suite 1370 for approximately 4,064 rentable square feet.\n"
        "Commencement Date September 1, 2026\n"
        "Lease Term Thirty-eight (38) months\n"
        "Lease Rate $45.00 per square foot per year NNN with 3% annual escalations.\n"
        "Operating Expenses Tenant shall pay its pro rata share of building expenses, which are estimated to be $24.44 per square foot for 2025.\n"
        "Parking Tenant shall be entitled to the non-exclusive use of 1 parking space per 600 rentable square feet.\n"
    )

    assert _should_ocr_docx_images(text) is False
    assert schedule[0] == {"start_month": 0, "end_month": 11, "rent_psf_annual": 42.0}
    assert schedule[1]["rent_psf_annual"] == 43.26
    assert schedule[-1]["end_month"] == 59


def test_regex_prefill_prefers_tia_per_sf_over_total_budget() -> None:
    text = (
        "Tenant Improvement Allowance (TIA): $35.00 per RSF.\n"
        "Tenant Improvement Budget: $240,000 total buildout allowance.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 35.0


def test_regex_prefill_does_not_treat_test_fit_or_opex_as_ti_allowance() -> None:
    text = (
        "TEST FIT: Landlord shall provide a test-fit allowance of $.12 per square foot, outside of the Tenant Improvement Allowance.\n"
        "Operating Expenses are estimated to be $14.30 per RSF.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") is None


def test_regex_prefill_does_not_treat_moving_furniture_or_signage_allowances_as_ti() -> None:
    text = (
        "Moving Allowance: Landlord shall reimburse up to $2.50 per RSF for relocation costs.\n"
        "Furniture Allowance: Landlord shall provide a $7.00 per RSF furniture allowance.\n"
        "Signage Allowance: Landlord shall provide $3,000 for exterior building signage.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") is None
    assert prefill.get("ti_allowance_total") is None


def test_regex_prefill_does_not_treat_turnkey_or_landlord_work_as_ti_allowance() -> None:
    text = (
        "Landlord Work: Premises shall be delivered on a turnkey basis pursuant to the Work Letter.\n"
        "Landlord shall complete the buildout and turnkey improvements at its sole cost.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") is None
    assert prefill.get("ti_allowance_total") is None


def test_regex_prefill_keeps_true_ti_allowance_when_other_allowances_are_present() -> None:
    text = (
        "Tenant Improvement Allowance: Landlord shall provide $22.00 per RSF.\n"
        "Moving Allowance: Landlord shall reimburse up to $2.50 per RSF for relocation costs.\n"
        "Furniture Allowance: $7.00 per RSF.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 22.0


def test_regex_prefill_prefers_counter_ti_allowance_and_supports_square_foot_unit() -> None:
    text = (
        "Tenant Improvements | Landlord to provide Tenant with a Tenant Improvement Allowance of $18.00 per square foot.\n"
        "Landlord shall provide Tenant an allowance not to exceed $20.00 per RSF from the Premises as-is condition.\n"
        "Landlord shall provide tenant an allowance up to $0.15/per RSF for a test-fit and one (1) revision.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 20.0


def test_regex_prefill_parses_ti_allowance_from_tenant_improvements_clause() -> None:
    text = (
        "Tenant Improvements: Premises is part of Landlord's spec suite program.\n"
        "Tenant shall be provided an allowance of $1.50 RSF for power to furniture.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 1.5


def test_regex_prefill_parses_ti_allowance_when_heading_and_body_are_merged() -> None:
    text = (
        "Premises: Portion of Suite 500, 4,165 RSF.\n"
        "Tenant ImprovementsLandlord will provide an allowance equal to $50.00 per square foot, inclusive of architectural and engineering fees.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 50.0


def test_regex_prefill_prefers_split_line_subtenant_allowance_over_tia_date_numbers() -> None:
    text = (
        "Rent Commencement Date | April 1, 2028\n"
        "Subtenant Allowance & Improvements\n"
        "$215.00 per RSF\n"
        "Sublandlord shall provide an allowance equal to the above for construction of improvements in the Premises (the \"TIA\").\n"
        "Any unused portion of the Subtenant Allowance remaining as of May 1, 2029 shall remain with Sublandlord.\n"
        "Test Fit Allowance | Sublandlord shall reimburse one test fit not to exceed $0.15/RSF.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_psf") == 215.0


def test_extract_lease_hints_parses_ti_allowance_when_heading_and_body_are_merged() -> None:
    text = (
        "Building: Domain Place\n"
        "Premises: Portion of Suite 500, 4,165 RSF.\n"
        "Lease Commencement: October 1, 2026.\n"
        "Lease Expiration: March 31, 2033.\n"
        "Tenant ImprovementsLandlord will provide an allowance equal to $50.00 per square foot, inclusive of architectural and engineering fees.\n"
    )
    hints = main._extract_lease_hints(text, "proposal.docx", "test-rid")
    assert hints.get("ti_allowance_psf") == 50.0


def test_regex_prefill_parses_total_ti_allowance_and_converts_to_psf() -> None:
    text = (
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Tenant Improvement Allowance: Owner shall provide Tenant with a Tenant Improvement Allowance of $50,000.00.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("ti_allowance_total") == 50000.0
    assert abs(float(prefill.get("ti_allowance_psf") or 0.0) - (50000.0 / 3947.0)) < 0.01


def test_extract_lease_hints_parses_total_ti_allowance_in_docx_style_text() -> None:
    text = (
        "Building: Lamar Central\n"
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Commencement Date: May 1, 2027.\n"
        "Lease Term: Sixty (60) months from the Commencement Date.\n"
        "Tenant Improvement Allowance: Owner shall provide Tenant with a Tenant Improvement Allowance, of $50,000.00.\n"
    )
    hints = main._extract_lease_hints(text, "proposal.docx", "test-rid")
    assert hints.get("ti_allowance_total") == 50000.0
    assert abs(float(hints.get("ti_allowance_psf") or 0.0) - (50000.0 / 3947.0)) < 0.01


def test_regex_prefill_ignores_timeline_dates_and_parses_ordinal_and_through_dates() -> None:
    text = (
        "LL Proposal: November 12, 2025\n"
        "Tenant Response December 23, 2025\n"
        "BUILDING: | Aspen Lake 2 - 10124 Lake Creek Pkwy | Aspen Lake 2 - 10124 Lake Creek Pkwy\n"
        "PREMISES: | 128,990 rentable square feet making up the entire Building.\n"
        "COMMENCEMENT DATE: | May 1st, 2028.\n"
        "LEASE TERM: | Through August, 30, 2033.\n"
        "BASE ANNUAL NET RENTAL RATE: | Months 1-12: Annual Base Rent: $29.69/RSF\n"
    )
    prefill = _regex_prefill(text)

    assert prefill.get("commencement") == "2028-05-01"
    assert prefill.get("expiration") == "2033-08-30"
    assert prefill.get("term_months") in {63, 64}
    assert prefill.get("name") == "Aspen Lake 2"


def test_regex_prefill_parses_annual_anniversary_increase_phrase() -> None:
    text = (
        "Premises: Suite 400 consisting of 4,949 RSF.\n"
        "Lease Commencement: December 1, 2026.\n"
        "Lease Term: One hundred twenty-seven (127) months.\n"
        "Base Rent: $33.00 per RSF NNN with 3% increases on the annual anniversary of the Lease Commencement.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 33.0}
    assert steps[1]["rate_psf_yr"] == 33.99
    assert steps[-1]["end"] == 126


def test_regex_prefill_overrides_flat_year_table_when_escalation_clause_is_explicit() -> None:
    text = (
        "Premises: The Premises will consist of approximately 51,743 RSF located in Buildings 100, 200, 300 and Amenity Building.\n"
        "Lease Commencement Date: August 1, 2026.\n"
        "Lease Term: One hundred twenty-eight (128) months.\n"
        "Base Rent: $47.50 per RSF per year NNN.\n"
        "Beginning in Month 13 after the Rent Commencement Date, the Base Rental Rate will increase by 3.0% annually.\n"
        "Lease Year 1 | 8/1/2026 | 7/31/2027 | 47.50\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 47.5}
    assert steps[1] == {"start": 12, "end": 23, "rate_psf_yr": 48.93}
    assert steps[-1] == {"start": 120, "end": 127, "rate_psf_yr": 63.84}


def test_regex_prefill_parses_label_value_escalation_with_percent_parenthetical() -> None:
    text = (
        "Premises: Suite 230 consisting of 3,947 RSF.\n"
        "Lease Commencement: May 1, 2027.\n"
        "Lease Expiration: April 30, 2032.\n"
        "Base Rent | $42.00/SF\n"
        "Annual Base Rent Escalation (%) | 3.00%\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("_rent_steps_source") == "base_rate_plus_escalation_regex"
    steps = prefill.get("rent_steps")
    assert isinstance(steps, list)
    assert steps[0] == {"start": 0, "end": 11, "rate_psf_yr": 42.0}
    assert steps[1]["rate_psf_yr"] == 43.26


def test_regex_prefill_prefers_parking_ratio_derived_count_over_small_inline_count() -> None:
    text = (
        "Premises: Suite 400 consisting of 4,949 RSF.\n"
        "Parking ratio per 1,000 RSF: 4.0\n"
        "Parking: on a must-take basis, four (4) parking spaces at $100 per month.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("parking_ratio_per_1000_rsf") == 4.0
    assert prefill.get("parking_spaces") == 20
    assert prefill.get("parking_cost_monthly_per_space") == 100.0


def test_regex_prefill_derives_parking_ratio_from_count_when_ratio_missing() -> None:
    text = (
        "Premises: Suite 400 consisting of 4,949 RSF.\n"
        "Parking: on a must-take, must-pay basis, four (4) parking spaces at $100 per month.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("parking_spaces") == 4
    assert round(float(prefill.get("parking_ratio_per_1000_rsf")), 4) == round((4 * 1000.0) / 4949.0, 4)
    assert prefill.get("parking_cost_monthly_per_space") == 100.0


def test_regex_prefill_parses_parking_ratio_per_every_1000_without_false_decimal_tail_count() -> None:
    text = (
        "Premises: Portion of Suite 440 consisting of 3,000 RSF.\n"
        "PARKING: Tenant's parking ratio shall be 1.32 parking spaces per every 1,000 rentable square feet of Lease space on a reserved basis.\n"
        "All parking in the building garage shall be $200 per space per month plus applicable sales tax.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("parking_ratio_per_1000_rsf") == 1.32
    assert prefill.get("parking_spaces") == 4
    assert prefill.get("parking_cost_monthly_per_space") == 200.0


def test_regex_prefill_handles_access_card_parking_and_nnn_with_base_year_reference() -> None:
    text = (
        "Premises: Suite 1550 consisting of 5,618 rentable square feet.\n"
        "Base Rental Rate: Months 1-7: $0.00 PSF + NNN; Months 8-12: $52.00 PSF + NNN.\n"
        "Operating Expenses: Tenant shall pay its pro rata share of actual NNN operating expenses during the term. "
        "Operating expenses for the year 2026 are estimated to be $20.90. "
        "Tenant shall have a cap on controllable opex of 6% per year, using a base year of 2027.\n"
        "PARKING: Tenant shall be provided non-reserved parking at a ratio of 2.0 access cards/1,000 RSF.\n"
        "Current charge for parking is $225/access card/month plus tax for non-reserved parking.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("opex_mode") == "nnn"
    assert prefill.get("base_opex_psf_yr") == 20.9
    assert prefill.get("free_rent_months") == 7
    assert prefill.get("parking_ratio_per_1000_rsf") == 2.0
    assert prefill.get("parking_spaces") == 11
    assert prefill.get("parking_cost_monthly_per_space") == 225.0


def test_regex_prefill_counterproposal_ignores_fit_plan_ti_and_derives_parking_from_600_rsf_ratio() -> None:
    text = (
        "Premises Suite 1370 for approximately 4,064 rentable square feet.\n"
        "Commencement Date September 1, 2026\n"
        "Lease Term Thirty-eight (38) months\n"
        "Rental Abatement Tenant shall have (4) months of rental abatement. Tenant shall be responsible for the operating expenses during this abatement period.\n"
        "Lease Rate $45.00 per square foot per year NNN with 3% annual escalations.\n"
        "Operating Expenses Tenant shall pay its pro rata share of building expenses, which are estimated to be $24.44 per square foot for 2025 (CAM $12.57, Taxes $11.53, Insurance $0.34).\n"
        "Initial Fit Plan Landlord shall provide an initial fit plan allowance of $0.15/RSF.\n"
        "Tenant Improvements Landlord will pay a tenant allowance of $5.00 per square foot, inclusive of architectural and engineering fees, as well as a four percent (4%) construction management fee on total costs.\n"
        "Parking Tenant shall be entitled to the non-exclusive use of 1 parking space per 600 rentable square feet. The current rate for unreserved parking is $250.00 per space plus tax.\n"
    )

    prefill = _regex_prefill(text)

    assert prefill.get("ti_allowance_psf") == 5.0
    assert prefill.get("base_opex_psf_yr") == 24.44
    assert round(float(prefill.get("parking_ratio_per_1000_rsf") or 0.0), 4) == round(1000.0 / 600.0, 4)
    assert prefill.get("parking_spaces") == 7
    assert prefill.get("parking_cost_monthly_per_space") == 250.0


def test_regex_prefill_free_rent_ignores_renewal_notice_months() -> None:
    text = (
        "TERM AND FREE RENT: One hundred twenty-seven (127) months, with seven (7) months base free rent.\n"
        "RENEWAL RIGHT: Tenant shall provide written notice no earlier than twelve (12) months nor later than nine (9) months prior to lease expiration.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("term_months") == 127
    assert prefill.get("free_rent_months") == 7


def test_regex_prefill_ignores_termination_fee_months_when_free_rent_is_present() -> None:
    text = (
        "Free Base Rent: Tenant shall receive four (4) months of abated base rent at the beginning of the lease term.\n"
        "Termination fee shall include unamortized costs (including free rent) and 7 months' gross rent.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("free_rent_months") == 4


def test_regex_prefill_ignores_holdover_months_and_prefers_primary_term() -> None:
    text = (
        "Lease Commencement Date | May 1, 2026.\n"
        "Lease Term | Three (3) years.\n"
        "Landlord Proposes a Thirty-Nine (39) month Lease Term.\n"
        "Holdover | In the event that Tenant possesses the Premises beyond the expiration of the Lease Term "
        "for the first three (3) months of such holdover, the Base Rent shall be one hundred twenty five percent (125%).\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("term_months") == 39
    assert prefill.get("expiration") == "2029-07-31"


def test_regex_prefill_handles_split_line_term_dates_and_ti_psf_without_false_ti_total() -> None:
    text = (
        "LL January 21, 2026\n"
        "COMMENCEMENT:\n"
        "September 1, 2026\n"
        "RENEWAL TERM:\n"
        "Ninety(90) months\n"
        "RENEWAL BASE RENT:\n"
        "$13.50NNN with 3.50% annual increases\n"
        "TENANT IMPROVEMENT ALLOWANCE:\n"
        "Landlord shall provide the Tenant a Tenant Improvement allowance equal to $1300 PSF for improvements.\n"
        "Tenant may have the ability to amortize an additional $7.00 PSF at 9% interest over the term of the lease.\n"
    )
    prefill = _regex_prefill(text)
    assert prefill.get("commencement") == "2026-09-01"
    assert prefill.get("term_months") == 90
    assert prefill.get("expiration") == "2034-02-28"
    assert prefill.get("rate_psf_yr") == 13.5
    assert prefill.get("ti_allowance_psf") == 13.0
    assert prefill.get("ti_allowance_total") is None
