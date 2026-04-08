from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import extraction.pipeline as extraction_pipeline
from docx import Document
from extraction.normalize import NormalizedDocument, PageData
from extraction.classify import classify_document
from extraction.regex import mine_candidates
from extraction.reconcile import reconcile
from extraction.schema import validate_canonical_extraction
from extraction.validate import validate_extraction
from services.input_normalizer import _dict_to_canonical
from models.canonical_lease import LeaseType


def _fixture_path(name: str) -> Path:
    return Path(__file__).parent / "fixtures" / "expected_json" / name


def test_schema_validation_fixture_passes() -> None:
    payload = json.loads(_fixture_path("valid_extraction.json").read_text(encoding="utf-8"))
    ok, errors = validate_canonical_extraction(payload)
    assert ok, f"schema should validate fixture, got: {errors}"


def test_rent_coverage_validation() -> None:
    extraction = {
        "document": {"doc_type": "lease", "doc_role": "prime_lease", "confidence": 0.9, "evidence_spans": []},
        "term": {
            "commencement_date": "2026-01-01",
            "expiration_date": "2026-12-31",
            "rent_commencement_date": "2026-01-01",
            "term_months": 12,
        },
        "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 1000},
        "rent_steps": [{"start_month": 0, "end_month": 11, "rate_psf_annual": 40.0}],
        "abatements": [],
        "opex": {"mode": "nnn", "base_psf_year_1": 10.0, "growth_rate": 0.03, "cues": ["nnn"]},
        "provenance": {},
        "review_tasks": [],
        "evidence": [{"source": "pdf_text_regex", "source_confidence": 0.9, "snippet": "NNN", "page": 1, "bbox": None}],
        "confidence": {"overall": 0.5, "status": "yellow", "export_allowed": True},
        "export_allowed": True,
    }
    result = validate_extraction(extraction, source_quality=0.9, reconcile_margin=0.5)
    assert result["export_allowed"] is True
    assert not any(t.get("issue_code") == "RENT_SCHEDULE_COVERAGE" for t in result["review_tasks"])


def test_term_constraint_within_tolerance() -> None:
    extraction = {
        "document": {"doc_type": "lease", "doc_role": "prime_lease", "confidence": 0.9, "evidence_spans": []},
        "term": {
            "commencement_date": "2026-01-01",
            "expiration_date": "2026-12-31",
            "rent_commencement_date": "2026-01-01",
            "term_months": 11,
        },
        "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 1000},
        "rent_steps": [{"start_month": 0, "end_month": 10, "rate_psf_annual": 40.0}],
        "abatements": [],
        "opex": {"mode": "nnn", "base_psf_year_1": 10.0, "growth_rate": 0.03, "cues": ["nnn"]},
        "provenance": {},
        "review_tasks": [],
        "evidence": [{"source": "pdf_text_regex", "source_confidence": 0.9, "snippet": "NNN", "page": 1, "bbox": None}],
        "confidence": {"overall": 0.5, "status": "yellow", "export_allowed": True},
        "export_allowed": True,
    }
    result = validate_extraction(extraction, source_quality=0.9, reconcile_margin=0.5)
    assert not any(t.get("issue_code") == "TERM_MISMATCH" for t in result["review_tasks"])


def test_opex_gating_when_nnn_cues_present_and_opex_missing() -> None:
    extraction = {
        "document": {"doc_type": "lease", "doc_role": "prime_lease", "confidence": 0.9, "evidence_spans": []},
        "term": {
            "commencement_date": "2026-01-01",
            "expiration_date": "2026-12-31",
            "rent_commencement_date": "2026-01-01",
            "term_months": 11,
        },
        "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 1000},
        "rent_steps": [{"start_month": 0, "end_month": 10, "rate_psf_annual": 40.0}],
        "abatements": [],
        "opex": {"mode": "", "base_psf_year_1": None, "growth_rate": None, "cues": []},
        "provenance": {},
        "review_tasks": [],
        "evidence": [{"source": "pdf_text_regex", "source_confidence": 0.9, "snippet": "Lease is NNN.", "page": 1, "bbox": None}],
        "confidence": {"overall": 0.5, "status": "yellow", "export_allowed": True},
        "export_allowed": True,
    }
    result = validate_extraction(extraction, source_quality=0.7, reconcile_margin=0.4)
    assert result["export_allowed"] is False
    blocker_codes = [t.get("issue_code") for t in result["review_tasks"] if t.get("severity") == "blocker"]
    assert "OPEX_NNN_INCOMPLETE" in blocker_codes


def test_pipeline_sanitizes_non_schema_review_task_severities() -> None:
    sanitized = extraction_pipeline._sanitize_review_tasks_for_schema(
        [
            {
                "field_path": "rent_steps",
                "severity": "low",
                "issue_code": "INCOMPLETE_RENT_SCHEDULE",
                "message": "Only the first rent row was captured.",
                "candidates": None,
                "evidence": None,
            }
        ]
    )

    assert sanitized == [
        {
            "field_path": "rent_steps",
            "severity": "warn",
            "issue_code": "INCOMPLETE_RENT_SCHEDULE",
            "message": "Only the first rent row was captured.",
            "candidates": [],
            "evidence": [],
        }
    ]


def test_arbitration_conflict_has_resolution_or_review_task() -> None:
    regex_candidates = {
        "term_months": [
            {"field": "term_months", "value": 60, "source": "pdf_text_regex", "source_confidence": 0.70, "snippet": "term 60 months", "page": 1, "bbox": None},
            {"field": "term_months", "value": 61, "source": "pdf_text_regex", "source_confidence": 0.69, "snippet": "term 61 months", "page": 2, "bbox": None},
        ]
    }
    rent_steps = [
        {"start_month": 0, "end_month": 59, "rate_psf_annual": 40.0, "source": "table_parser", "source_confidence": 0.71, "snippet": "Year 1-5", "page": 3, "bbox": None},
        {"start_month": 0, "end_month": 59, "rate_psf_annual": 41.0, "source": "table_parser", "source_confidence": 0.70, "snippet": "Year 1-5", "page": 4, "bbox": None},
    ]

    reconciled = reconcile(regex_candidates=regex_candidates, rent_step_candidates=rent_steps, llm_output=None)
    resolved_term = (reconciled.get("resolved") or {}).get("term", {}).get("term_months")
    provenance = (reconciled.get("provenance") or {}).get("term.term_months", [])

    assert resolved_term in {60, 61} or resolved_term is None
    assert provenance, "conflict resolution must include evidence provenance"


def test_regex_detects_full_service_gross_opex_mode() -> None:
    normalized = NormalizedDocument(
        sha256="x",
        filename="sample.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Lease Type: Full Service Gross Lease. Tenant will not pay separate CAM.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Lease Type: Full Service Gross Lease.",
    )
    candidates = mine_candidates(normalized)
    values = [str(c.get("value") or "") for c in candidates.get("opex_mode", [])]
    assert "full_service" in values


def test_regex_opex_mode_prefers_nnn_when_base_year_is_only_cap_reference() -> None:
    normalized = NormalizedDocument(
        sha256="n",
        filename="sample.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text=(
                    "Operating Expenses: Tenant shall pay its pro rata share of actual NNN operating expenses. "
                    "Operating expenses for 2026 are estimated to be $20.90. "
                    "Tenant shall have a cap on controllable opex using a base year of 2027."
                ),
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="NNN operating expenses with base year cap reference",
    )
    candidates = mine_candidates(normalized)
    values = [str(c.get("value") or "") for c in candidates.get("opex_mode", [])]
    assert "nnn" in values


def test_reconcile_detects_phase_in_without_abatement() -> None:
    normalized = NormalizedDocument(
        sha256="phasein",
        filename="phasein.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Phase-In occupancy schedule: Months 1-12 at 12,500 RSF, Months 13-24 at 21,000 RSF.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Phase-In occupancy schedule",
    )
    candidates = mine_candidates(normalized)
    reconciled = reconcile(regex_candidates=candidates, rent_step_candidates=[], llm_output=None)
    resolved = reconciled.get("resolved") or {}
    analysis = resolved.get("abatement_analysis") or {}
    assert analysis.get("classification") == "phase_in"
    assert analysis.get("phase_in_detected") is True
    assert (resolved.get("abatements") or []) == []


def test_reconcile_detects_gross_vs_base_abatement_scope() -> None:
    gross_doc = NormalizedDocument(
        sha256="gross",
        filename="gross.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Tenant shall receive gross rent abatement for months 1-4, including base rent and operating expenses.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="gross rent abatement",
    )
    base_doc = NormalizedDocument(
        sha256="base",
        filename="base.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Landlord grants base rent only abatement during months 1-2.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="base rent only abatement",
    )

    gross_reconciled = reconcile(regex_candidates=mine_candidates(gross_doc), rent_step_candidates=[], llm_output=None)
    base_reconciled = reconcile(regex_candidates=mine_candidates(base_doc), rent_step_candidates=[], llm_output=None)

    gross_abatements = (gross_reconciled.get("resolved") or {}).get("abatements") or []
    base_abatements = (base_reconciled.get("resolved") or {}).get("abatements") or []
    assert gross_abatements and gross_abatements[0].get("scope") == "gross_rent"
    assert base_abatements and base_abatements[0].get("scope") == "base_rent_only"


def test_dict_normalizer_maps_full_service_variants() -> None:
    lease = _dict_to_canonical(
        {
            "scenario_name": "Option 1",
            "term_months": 60,
            "rsf": 1000,
            "commencement_date": "2026-01-01",
            "expiration_date": "2031-01-01",
            "lease_type": "Full Service Gross Lease",
            "rent_schedule": [{"start_month": 0, "end_month": 59, "rent_psf_annual": 40}],
        }
    )
    assert lease.lease_type == LeaseType.FULL_SERVICE


def test_reconcile_extracts_ti_parking_and_rights_into_structured_output() -> None:
    normalized = NormalizedDocument(
        sha256="cre-rich",
        filename="proposal.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text=(
                    "Tenant Improvement Allowance: $80.00 per RSF.\n"
                    "Parking: 3.5 / 1,000 RSF at $225 per space per month.\n"
                    "Renewal Option: 1 x 5 year option at FMV.\n"
                    "The first 6 months of free rent shall be abated."
                ),
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="rich proposal text",
    )
    candidates = mine_candidates(normalized)
    reconciled = reconcile(regex_candidates=candidates, rent_step_candidates=[], llm_output=None)
    resolved = reconciled.get("resolved") or {}

    assert (resolved.get("tenant_improvements") or {}).get("ti_allowance_psf") == 80.0
    assert (resolved.get("parking") or {}).get("ratio_per_1000_rsf") == 3.5
    assert (resolved.get("parking") or {}).get("rate_monthly_per_space") == 225.0
    assert (resolved.get("rights_options") or {}).get("renewal_option") == "Renewal Option: 1 x 5 year option at FMV."
    assert (resolved.get("concessions") or {}).get("free_rent_months") == 6


def test_validate_extraction_reports_missing_information_for_model_ready_fields() -> None:
    extraction = {
        "document": {"doc_type": "proposal", "doc_role": "proposal", "confidence": 0.9, "evidence_spans": []},
        "term": {
            "commencement_date": None,
            "expiration_date": None,
            "rent_commencement_date": None,
            "term_months": 60,
        },
        "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 1000},
        "rent_steps": [{"start_month": 0, "end_month": 59, "rate_psf_annual": 40.0}],
        "abatements": [],
        "abatement_analysis": {"classification": "none", "phase_in_detected": False, "phase_in_confidence": 0.0, "scope": None},
        "concessions": {},
        "tenant_improvements": {},
        "parking": {},
        "rights_options": {},
        "opex": {"mode": "nnn", "base_psf_year_1": 10.0, "growth_rate": 0.03, "cues": ["nnn"]},
        "provenance": {},
        "review_tasks": [],
        "evidence": [{"source": "pdf_text_regex", "source_confidence": 0.9, "snippet": "NNN", "page": 1, "bbox": None}],
        "confidence": {"overall": 0.5, "status": "yellow", "export_allowed": True},
        "export_allowed": True,
    }
    result = validate_extraction(extraction, source_quality=0.8, reconcile_margin=0.5)
    missing = set(result.get("missing_information") or [])
    assert "term.commencement_date" in missing
    assert "tenant_improvements.ti_allowance_psf" in missing
    assert "parking.ratio_per_1000_rsf" in missing
    issue_codes = {t.get("issue_code") for t in result.get("review_tasks") or []}
    assert "MISSING_TENANT_IMPROVEMENTS_TI_ALLOWANCE_PSF" in issue_codes


def test_reconcile_derives_ti_total_from_ti_psf_and_rsf() -> None:
    normalized = NormalizedDocument(
        sha256="ti-derive",
        filename="proposal.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Premises: 5,000 RSF. Tenant Improvement Allowance: $40.00 per RSF.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Premises and TI allowance",
    )

    reconciled = reconcile(regex_candidates=mine_candidates(normalized), rent_step_candidates=[], llm_output=None)
    tenant_improvements = (reconciled.get("resolved") or {}).get("tenant_improvements") or {}

    assert tenant_improvements.get("ti_allowance_psf") == 40.0
    assert tenant_improvements.get("ti_allowance_total") == 200000.0


def test_validate_extraction_flags_concession_and_ti_incompleteness() -> None:
    extraction = {
        "document": {"doc_type": "proposal", "doc_role": "proposal", "confidence": 0.9, "evidence_spans": []},
        "term": {
            "commencement_date": "2026-01-01",
            "expiration_date": "2030-12-31",
            "rent_commencement_date": "2026-01-01",
            "term_months": 60,
        },
        "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 0},
        "rent_steps": [{"start_month": 0, "end_month": 59, "rate_psf_annual": 40.0}],
        "abatements": [],
        "abatement_analysis": {"classification": "none", "phase_in_detected": False, "phase_in_confidence": 0.0, "scope": None},
        "concessions": {},
        "tenant_improvements": {"ti_allowance_total": 50000.0},
        "parking": {},
        "rights_options": {},
        "opex": {"mode": "nnn", "base_psf_year_1": 10.0, "growth_rate": 0.03, "cues": ["nnn"]},
        "provenance": {},
        "review_tasks": [],
        "evidence": [
            {
                "source": "pdf_text_regex",
                "source_confidence": 0.9,
                "snippet": (
                    "Tenant Improvement Allowance: $50,000 total. "
                    "Moving Allowance: $2.50 per RSF. "
                    "Tenant shall receive free rent during the initial term."
                ),
                "page": 1,
                "bbox": None,
            }
        ],
        "confidence": {"overall": 0.5, "status": "yellow", "export_allowed": True},
        "export_allowed": True,
    }

    result = validate_extraction(extraction, source_quality=0.8, reconcile_margin=0.5)
    issue_codes = {t.get("issue_code") for t in result.get("review_tasks") or []}
    assert "TI_ALLOWANCE_NORMALIZATION_INCOMPLETE" in issue_codes
    assert "FREE_RENT_INCOMPLETE" in issue_codes
    assert "MOVING_ALLOWANCE_DETECTED" in issue_codes


def test_rule_classifier_detects_flyer_and_floorplan() -> None:
    flyer = NormalizedDocument(
        sha256="flyer",
        filename="domain-flyer.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Marketing Flyer\nProperty Highlights\nAvailable suites and asking rent listed below.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Marketing flyer for available suites.",
    )
    floorplan = NormalizedDocument(
        sha256="floorplan",
        filename="tower-floorplan.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Floor Plan / Stacking Plan\nLevel 15 layout with Suite 1550 and Suite 1560.",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Floor plan and stacking plan details.",
    )

    assert classify_document(flyer).get("doc_type") == "flyer"
    assert classify_document(floorplan).get("doc_type") == "floorplan"


def test_run_extraction_pipeline_surfaces_rejected_rent_rows_as_review_tasks(monkeypatch) -> None:
    normalized = NormalizedDocument(
        sha256="pipeline-rent-review",
        filename="lease.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text="Lease Year 1 Monthly Rent $25,000",
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text="Lease Year 1 Monthly Rent $25,000",
    )

    monkeypatch.setattr(extraction_pipeline, "normalize_document", lambda *args, **kwargs: normalized)
    monkeypatch.setattr(
        extraction_pipeline,
        "extract_rent_step_candidates_with_review",
        lambda *args, **kwargs: (
            [],
            [
                {
                    "field_path": "rent_steps",
                    "severity": "warn",
                    "issue_code": "RENT_ROW_REJECTED_MONTHLY_TOTAL",
                    "message": "Row looked like monthly total rent, not annual PSF rent.",
                    "candidates": [
                        {
                            "row_text": "Lease Year 1 Monthly Rent $25,000",
                            "category": "monthly_total",
                            "reason": "Row looked like monthly total rent, not annual PSF rent.",
                            "source": "text_line_regex",
                        }
                    ],
                    "recommended_value": None,
                    "evidence": [
                        {
                            "page": 1,
                            "snippet": "Lease Year 1 Monthly Rent $25,000",
                            "bbox": None,
                            "source": "text_line_regex",
                            "source_confidence": 0.9,
                        }
                    ],
                    "metadata": {
                        "row_text": "Lease Year 1 Monthly Rent $25,000",
                        "rejection_reason": "Row looked like monthly total rent, not annual PSF rent.",
                        "rejection_category": "monthly_total",
                        "base_issue_code": "RENT_ROW_REJECTED",
                    },
                }
            ],
        ),
    )
    monkeypatch.setattr(extraction_pipeline, "mine_candidates", lambda *args, **kwargs: {})
    monkeypatch.setattr(extraction_pipeline, "retrieve_section_snippets", lambda *args, **kwargs: {})
    monkeypatch.setattr(
        extraction_pipeline,
        "classify_document",
        lambda *args, **kwargs: {"doc_type": "lease", "doc_role": "prime_lease", "confidence": 0.9, "evidence_spans": []},
    )
    monkeypatch.setattr(extraction_pipeline, "structured_extract", lambda *args, **kwargs: {"review_tasks": []})
    monkeypatch.setattr(
        extraction_pipeline,
        "reconcile",
        lambda *args, **kwargs: {
            "resolved": {
                "term": {"commencement_date": "2026-01-01", "expiration_date": "2026-12-31", "term_months": 12},
                "premises": {"building_name": "A", "suite": "100", "floor": None, "address": "A", "rsf": 1000},
                "rent_steps": [],
                "abatements": [],
                "abatement_analysis": {"classification": "none", "phase_in_detected": False, "phase_in_confidence": 0.0, "scope": None},
                "concessions": {},
                "tenant_improvements": {},
                "parking": {},
                "rights_options": {},
                "opex": {"mode": "nnn", "base_psf_year_1": 10.0, "growth_rate": 0.03, "cues": ["nnn"]},
            },
            "provenance": {},
            "reconcile_margin": 0.5,
            "solver_debug": {},
            "degraded": False,
        },
    )

    result = extraction_pipeline.run_extraction_pipeline(
        file_bytes=b"%PDF-1.4 test",
        filename="lease.pdf",
        content_type="application/pdf",
    )

    rent_review = next(task for task in result.get("review_tasks") or [] if task.get("issue_code") == "RENT_ROW_REJECTED_MONTHLY_TOTAL")
    assert rent_review.get("metadata", {}).get("rejection_category") == "monthly_total"
    assert "Monthly Rent" in rent_review.get("metadata", {}).get("row_text", "")


def test_reconcile_parses_distributed_abatements_and_parking_tie() -> None:
    normalized = NormalizedDocument(
        sha256="distributed-abatement",
        filename="proposal.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text=(
                    "Lease Term: 72 months.\n"
                    "Landlord shall provide three (3) months of base rent abatement and an additional four (4) months "
                    "of base rent abatement allocated as follows: two (2) months in the beginning of lease year 3 and "
                    "two (2) months in the beginning of lease year 5.\n"
                    "Parking costs shall be abated during the Abatement Period.\n"
                ),
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text=(
            "Lease Term: 72 months. "
            "Landlord shall provide three (3) months of base rent abatement and an additional four (4) months "
            "of base rent abatement allocated as follows: two (2) months in the beginning of lease year 3 and "
            "two (2) months in the beginning of lease year 5. "
            "Parking costs shall be abated during the Abatement Period."
        ),
    )

    reconciled = reconcile(
        regex_candidates=mine_candidates(normalized),
        rent_step_candidates=[],
        llm_output=None,
        full_text=normalized.full_text,
    )
    resolved = reconciled.get("resolved") or {}

    assert (resolved.get("concessions") or {}).get("free_rent_months") == 7
    assert (resolved.get("abatements") or []) == [
        {"start_month": 0, "end_month": 2, "scope": "base_rent_only", "classification": "rent_abatement"},
        {"start_month": 24, "end_month": 25, "scope": "base_rent_only", "classification": "rent_abatement"},
        {"start_month": 48, "end_month": 49, "scope": "base_rent_only", "classification": "rent_abatement"},
    ]
    assert (resolved.get("parking_abatements") or []) == [
        {"start_month": 0, "end_month": 2},
        {"start_month": 24, "end_month": 25},
        {"start_month": 48, "end_month": 49},
    ]


def test_run_extraction_pipeline_flags_unresolved_parking_abatement_period(monkeypatch) -> None:
    normalized = NormalizedDocument(
        sha256="parking-unclear",
        filename="lease.pdf",
        content_type="application/pdf",
        pages=[
            PageData(
                page_number=1,
                text=(
                    "Parking costs shall be abated during the Abatement Period.\n"
                    "Tenant shall receive gross rent abatement, as otherwise set forth herein.\n"
                ),
                words=[],
                table_regions=[],
                needs_ocr=False,
            )
        ],
        full_text=(
            "Parking costs shall be abated during the Abatement Period. "
            "Tenant shall receive gross rent abatement, as otherwise set forth herein."
        ),
    )

    monkeypatch.setattr(extraction_pipeline, "normalize_document", lambda *args, **kwargs: normalized)
    monkeypatch.setattr(extraction_pipeline, "extract_rent_step_candidates_with_review", lambda *args, **kwargs: ([], []))
    monkeypatch.setattr(extraction_pipeline, "retrieve_section_snippets", lambda *args, **kwargs: {})
    monkeypatch.setattr(
        extraction_pipeline,
        "classify_document",
        lambda *args, **kwargs: {"doc_type": "lease", "doc_role": "prime_lease", "confidence": 0.9, "evidence_spans": []},
    )
    monkeypatch.setattr(extraction_pipeline, "structured_extract", lambda *args, **kwargs: {"review_tasks": []})

    result = extraction_pipeline.run_extraction_pipeline(
        file_bytes=b"%PDF-1.4 parking unclear",
        filename="lease.pdf",
        content_type="application/pdf",
    )

    issue_codes = {task.get("issue_code") for task in result.get("review_tasks") or []}
    assert "FREE_RENT_INCOMPLETE" in issue_codes
    assert "PARKING_ABATEMENT_PERIOD_UNRESOLVED" in issue_codes


def test_run_extraction_pipeline_docx_proposal_uses_legacy_hint_fallback(monkeypatch) -> None:
    doc = Document()
    for line in (
        "REVISED Office Lease Proposal for Sprinklr (Tenant).",
        "Building: 1300 East 5th, located at 1300 E. 5th Street, Austin, Texas 78702.",
        "Premises: approximately 12,153 rentable square feet located on the 4th floor.",
        "Commencement Date: no later than December 1, 2026.",
        "Lease Term: Sixty-six (66) Months from the Commencement Date.",
        "Base Rental Rate: $43.50/RSF NNN. Beginning in Month 13, the Base Rent will be subject to annual escalations of 3.0%.",
        "Rental Abatement: So long as Tenant is not in default, Base Rent and Building Operating Expenses and Real Estate Taxes will be abated for the initial three (3) months of the Lease Term, followed by three (3) months of Base Rent abatement. During the Base Rent abatement period, Tenant will be responsible for Building Operating Expenses and Real Estate Taxes.",
        "In addition to the turn-key delivery, Landlord will provide an allowance equal to $5.00 per RSF that Tenant can use towards moving costs, FF&E, security, and data & cabling.",
        "Operating Expenses and Real Estate Taxes: 2026 estimated Operating Expenses and Real Estate Taxes are $20.06 per RSF.",
        "Parking: Tenant will have access to a parking ratio of 2.7 spaces per 1,000 RSF. There will be a charge of $185 per month per space for unreserved spaces. Tenant will have abated parking during the initial twenty-four (24) months of the Lease Term.",
        "Test-Fit: Landlord will provide a test-fit allowance equal to $0.15 per RSF of the Premises for Tenant's preferred architect for space planning.",
    ):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    monkeypatch.setattr(extraction_pipeline, "structured_extract", lambda *args, **kwargs: {"review_tasks": []})

    result = extraction_pipeline.run_extraction_pipeline(
        file_bytes=buf.getvalue(),
        filename="1300 E 5th_Sprinklr_LL_REVISED_4.6.26(5_Year).docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result["document"]["doc_type"] == "proposal"
    assert result["term"]["term_months"] == 66
    assert result["term"]["commencement_date"] == "2026-12-01"
    assert result["term"]["expiration_date"] == "2032-05-31"
    assert result["tenant_improvements"]["ti_allowance_psf"] == 5.0
    assert result["parking"]["ratio_per_1000_rsf"] == 2.7
    assert result["parking"]["rate_monthly_per_space"] == 185.0
    assert [(item["start_month"], item["end_month"], item["scope"]) for item in result["abatements"]] == [
        (0, 2, "gross_rent"),
        (3, 5, "base_rent_only"),
    ]
    assert [(item["start_month"], item["end_month"]) for item in result["parking_abatements"]] == [(0, 23)]
