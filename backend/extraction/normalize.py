from __future__ import annotations

import hashlib
import io
import os
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

from pypdf import PdfReader
from scenario_extract import extract_text_from_word


@dataclass
class WordToken:
    text: str
    page: int
    bbox: tuple[float, float, float, float] | None
    source: str


@dataclass
class TableRegion:
    page: int
    bbox: tuple[float, float, float, float] | None
    rows: list[list[str]] = field(default_factory=list)
    source: str = ""


@dataclass
class PageData:
    page_number: int
    text: str
    words: list[WordToken] = field(default_factory=list)
    table_regions: list[TableRegion] = field(default_factory=list)
    needs_ocr: bool = False


@dataclass
class NormalizedDocument:
    sha256: str
    filename: str
    content_type: str
    pages: list[PageData]
    full_text: str


_MIN_ALNUM_FOR_NATIVE_TEXT = 50
PDF_NORMALIZE_MAX_PAGES = max(1, int(os.environ.get("PDF_NORMALIZE_MAX_PAGES", "40")))
PDFPLUMBER_MAX_FILE_BYTES = max(1, int(os.environ.get("PDFPLUMBER_MAX_FILE_BYTES", str(8 * 1024 * 1024))))


def _is_weak_text(text: str, words: list[WordToken]) -> bool:
    alnum_count = len(re.findall(r"[A-Za-z0-9]", text or ""))
    return alnum_count < _MIN_ALNUM_FOR_NATIVE_TEXT or len(words) == 0


def _normalize_pdf_with_pymupdf(file_bytes: bytes, *, max_pages: int) -> list[PageData]:
    try:
        import fitz  # type: ignore
    except Exception:
        return []

    pages: list[PageData] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:  # type: ignore[arg-type]
        for i, page in enumerate(doc, start=1):
            if i > max_pages:
                break
            text = page.get_text("text", sort=True) or ""
            words_raw = page.get_text("words") or []
            words: list[WordToken] = []
            for entry in words_raw:
                if not entry or len(entry) < 5:
                    continue
                x0, y0, x1, y1, token = entry[:5]
                token_text = str(token or "").strip()
                if not token_text:
                    continue
                words.append(
                    WordToken(
                        text=token_text,
                        page=i,
                        bbox=(float(x0), float(y0), float(x1), float(y1)),
                        source="pymupdf_words",
                    )
                )
            pages.append(PageData(page_number=i, text=text, words=words, needs_ocr=_is_weak_text(text, words)))
    return pages


def _merge_pdfplumber_data(file_bytes: bytes, pages: list[PageData], *, max_pages: int) -> list[PageData]:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        return pages

    by_index: dict[int, PageData] = {p.page_number: p for p in pages}
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            if i > max_pages:
                break
            existing = by_index.get(i)
            if existing is None:
                existing = PageData(page_number=i, text="")
                by_index[i] = existing

            try:
                words = page.extract_words(keep_blank_chars=False, x_tolerance=1.5, y_tolerance=2.0) or []
            except Exception:
                words = []
            if words:
                existing.words.extend(
                    WordToken(
                        text=str(w.get("text") or "").strip(),
                        page=i,
                        bbox=(
                            float(w.get("x0") or 0.0),
                            float(w.get("top") or 0.0),
                            float(w.get("x1") or 0.0),
                            float(w.get("bottom") or 0.0),
                        ),
                        source="pdfplumber_words",
                    )
                    for w in words
                    if str(w.get("text") or "").strip()
                )

            try:
                found_tables = page.find_tables() or []
            except Exception:
                found_tables = []

            for tbl in found_tables:
                bbox = None
                try:
                    bbox_raw = getattr(tbl, "bbox", None)
                    if bbox_raw and len(bbox_raw) == 4:
                        bbox = tuple(float(x) for x in bbox_raw)  # type: ignore[assignment]
                except Exception:
                    bbox = None
                rows = []
                try:
                    extracted = tbl.extract() or []
                    rows = [[str(cell or "").strip() for cell in row] for row in extracted if isinstance(row, list)]
                except Exception:
                    rows = []
                existing.table_regions.append(TableRegion(page=i, bbox=bbox, rows=rows, source="pdfplumber_table"))

            if not existing.text.strip():
                try:
                    existing.text = page.extract_text() or ""
                except Exception:
                    existing.text = ""
            existing.needs_ocr = _is_weak_text(existing.text, existing.words)

    return [by_index[i] for i in sorted(by_index)]


def _normalize_pdf_with_pypdf(file_bytes: bytes, *, max_pages: int) -> list[PageData]:
    pages: list[PageData] = []
    reader = PdfReader(io.BytesIO(file_bytes))
    for i, page in enumerate(reader.pages, start=1):
        if i > max_pages:
            break
        text = page.extract_text() or ""
        pages.append(PageData(page_number=i, text=text, words=[], table_regions=[], needs_ocr=_is_weak_text(text, [])))
    return pages


def _normalize_docx(file_bytes: bytes) -> list[PageData]:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return [PageData(page_number=1, text="", words=[], table_regions=[], needs_ocr=False)]

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return [PageData(page_number=1, text="", words=[], table_regions=[], needs_ocr=False)]

    chunks: list[str] = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            chunks.append(txt)
    for tbl in doc.tables:
        for row in tbl.rows:
            values = [str(c.text or "").strip() for c in row.cells if str(c.text or "").strip()]
            if values:
                chunks.append(" | ".join(values))

    text = "\n".join(chunks)
    words = [WordToken(text=t, page=1, bbox=None, source="docx_token") for t in re.findall(r"\S+", text)]
    return [PageData(page_number=1, text=text, words=words, table_regions=[], needs_ocr=False)]


def _normalize_word(file_bytes: bytes, filename: str, content_type: str = "") -> list[PageData]:
    lower_name = (filename or "").lower()
    looks_like_docx = (
        lower_name.endswith(".docx")
        or "wordprocessingml" in (content_type or "").lower()
        or file_bytes.startswith(b"PK")
    )
    try:
        text, source = extract_text_from_word(BytesIO(file_bytes), filename=filename)
        token_source = f"{source}_token"
    except Exception:
        if looks_like_docx:
            return _normalize_docx(file_bytes)
        return [PageData(page_number=1, text="", words=[], table_regions=[], needs_ocr=False)]
    words = [WordToken(text=t, page=1, bbox=None, source=token_source) for t in re.findall(r"\S+", text)]
    return [PageData(page_number=1, text=text, words=words, table_regions=[], needs_ocr=False)]


def normalize_document(file_bytes: bytes, filename: str, content_type: str = "") -> NormalizedDocument:
    sha = hashlib.sha256(file_bytes).hexdigest()
    lower_name = (filename or "").lower()
    is_pdf = lower_name.endswith(".pdf") or "pdf" in (content_type or "").lower()

    if is_pdf:
        max_pages = max(1, PDF_NORMALIZE_MAX_PAGES)
        pages = _normalize_pdf_with_pymupdf(file_bytes, max_pages=max_pages)
        if not pages:
            pages = _normalize_pdf_with_pypdf(file_bytes, max_pages=max_pages)
        if len(file_bytes) <= PDFPLUMBER_MAX_FILE_BYTES:
            pages = _merge_pdfplumber_data(file_bytes, pages, max_pages=max_pages)
    else:
        pages = _normalize_word(file_bytes, filename, content_type)

    full_text = "\n\n".join((p.text or "") for p in pages)
    return NormalizedDocument(
        sha256=sha,
        filename=filename or "",
        content_type=content_type or "",
        pages=pages,
        full_text=full_text,
    )
