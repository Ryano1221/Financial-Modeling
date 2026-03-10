"""
Extract text from PDF (pypdf + OCR fallback) or DOCX (python-docx),
then run LLM to produce ExtractionResponse (Scenario + confidence + warnings).
User must always review before running financial model.
"""
from __future__ import annotations

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import time
import zipfile
import gc
from calendar import monthrange
from datetime import date, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, BinaryIO
import xml.etree.ElementTree as ET

from models import ExtractionResponse, Scenario, RentStep, OpexMode

logger = logging.getLogger(__name__)

# Early exit: skip OCR if PDF text is at least this long and high quality
OCR_EARLY_EXIT_MIN_CHARS = 1500
# When OCR runs, only process this many pages (default)
DEFAULT_OCR_PAGES = 3
# Cap OCR memory by limiting PDF size and render DPI.
OCR_MAX_PDF_BYTES = int(os.environ.get("OCR_MAX_PDF_BYTES", str(8 * 1024 * 1024)))
OCR_RENDER_DPI = int(os.environ.get("OCR_RENDER_DPI", "120"))
# DOCX image OCR controls (for embedded rent schedule screenshots/tables).
DOCX_OCR_MAX_IMAGES = int(os.environ.get("DOCX_OCR_MAX_IMAGES", "2"))
DOCX_OCR_MIN_IMAGE_AREA = int(os.environ.get("DOCX_OCR_MIN_IMAGE_AREA", str(200 * 120)))
DOCX_OCR_MAX_TOTAL_IMAGE_BYTES = int(os.environ.get("DOCX_OCR_MAX_TOTAL_IMAGE_BYTES", str(6 * 1024 * 1024)))
# Snippet pack cap
SNIPPET_PACK_MAX_CHARS = 20_000
# Fallback truncation when snippet pack too small
FALLBACK_HEAD_CHARS = 20_000
FALLBACK_TAIL_CHARS = 20_000
# Lines of context around keyword matches for snippet pack
SNIPPET_CONTEXT_LINES = 24
DOC_TEXT_COMMAND_TIMEOUT_SECONDS = 30
MIN_DOC_TEXT_CHARS = 80

FOCUS_KEYWORDS = [
    "rent", "base rent", "rental rate", "term", "commencement", "expiration",
    "operating expenses", "CAM", "NNN", "base year", "TI", "tenant improvement",
    "allowance", "abatement", "free rent", "parking", "option", "renewal", "termination",
    "full service", "full service gross", "gross lease", "modified gross", "expense stop",
    "premises", "suite", "building", "property", "address",
    "rentable square feet", "rentable area", "square feet", "rsf",
]
SNIPPET_MIN_CHARS = 500  # if snippet pack smaller than this, use fallback truncation

_NNN_OPEX_CUES = (
    " nnn ",
    "triple net",
    "triple-net",
    "n.n.n",
    "absolute nnn",
    "net lease",
)
_BASE_YEAR_OPEX_CUES = (
    "base year",
    "base-year",
    "expense stop",
    "gross with stop",
    "modified gross",
    "mod gross",
)
_FULL_SERVICE_OPEX_CUES = (
    "full service gross",
    "full-service gross",
    "full service lease",
    "full-service lease",
    "full service",
    "full-service",
    "gross lease",
    "fsg",
)


def _detect_opex_mode_from_text(raw_text: str) -> str | None:
    low = f" {(raw_text or '').lower()} "
    if any(k in low for k in _BASE_YEAR_OPEX_CUES):
        return "base_year"
    # Word-boundary NNN detection catches punctuation variants like ("NNN"), N.N.N., etc.
    if re.search(r"(?i)\b(?:n\.?\s*n\.?\s*n\.?|nnn)\b", raw_text or ""):
        return "nnn"
    if any(k in low for k in _FULL_SERVICE_OPEX_CUES):
        return "full_service"
    if any(k in low for k in _NNN_OPEX_CUES) and "not nnn" not in low and "no nnn" not in low:
        return "nnn"
    if re.search(r"(?i)\b(?:lease|expense|rent)\s*(?:type|structure|mode)\b[^\n]{0,40}\bgross\b", raw_text or ""):
        return "full_service"
    return None


def extract_text_from_pdf(file: BinaryIO) -> str:
    """Extract raw text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf required: pip install pypdf")
    reader = PdfReader(file)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts) if parts else ""


def _text_looks_high_quality(text: str) -> bool:
    """True if text has reasonable alphanumeric ratio (not scanned gibberish)."""
    if not text or len(text.strip()) < 100:
        return False
    alnum = sum(1 for c in text if c.isalnum() or c.isspace())
    return (alnum / len(text)) >= 0.6


def text_quality_requires_ocr(text: str) -> bool:
    """
    Return True if extracted text quality is poor and OCR should be run automatically.
    Uses: text length, alnum ratio, replacement-char count, short-line ratio.
    """
    t = (text or "").strip()
    text_len = len(t)
    if text_len == 0:
        return True
    non_ws = [c for c in t if not c.isspace()]
    non_ws_count = len(non_ws)
    alnum_count = sum(1 for c in non_ws if c.isalnum())
    alnum_ratio = alnum_count / max(1, non_ws_count)
    bad_char_count = t.count("\uFFFD")  # replacement character
    lines = [ln for ln in t.splitlines() if ln.strip()]
    short_lines = sum(1 for ln in lines if len(ln) < 20)
    short_line_ratio = short_lines / max(1, len(lines))

    if text_len < 1200:
        return True
    if alnum_ratio < 0.40:
        return True
    if bad_char_count > 5:
        return True
    if short_line_ratio > 0.60:
        return True
    return False


def _run_ocr_on_pdf(file: BinaryIO, max_pages: int = DEFAULT_OCR_PAGES) -> str:
    """Run OCR on first max_pages of PDF using pdf2image + pytesseract."""
    try:
        import pdf2image
        import pytesseract
    except ImportError as e:
        raise ImportError(
            "OCR requires pdf2image and pytesseract. "
            "Install: pip install pdf2image pytesseract. "
            "Also install system deps: poppler (e.g. brew install poppler) and tesseract (e.g. brew install tesseract)."
        ) from e
    t0 = time.perf_counter()
    tmp_pdf_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(prefix="ocr-input-", suffix=".pdf", delete=False) as tmp:
            file.seek(0)
            shutil.copyfileobj(file, tmp)
            tmp_pdf_path = tmp.name
    except Exception as e:
        raise ValueError(f"OCR failed (is poppler installed?): {e}") from e

    try:
        file_size = int(Path(tmp_pdf_path).stat().st_size) if tmp_pdf_path else 0
    except Exception:
        file_size = 0
    if file_size > OCR_MAX_PDF_BYTES:
        logger.warning(
            "[extract] OCR skipped: pdf bytes=%d exceeds OCR_MAX_PDF_BYTES=%d",
            file_size,
            OCR_MAX_PDF_BYTES,
        )
        if tmp_pdf_path:
            try:
                os.unlink(tmp_pdf_path)
            except Exception:
                pass
        return ""

    # Process one page at a time to avoid holding many rendered images in memory.
    total_pages = max_pages
    try:
        info = pdf2image.pdfinfo_from_path(tmp_pdf_path)
        detected_pages = int(info.get("Pages", max_pages))
        total_pages = max(0, min(max_pages, detected_pages))
    except Exception:
        # If page count probing fails, fall back to requested page limit.
        total_pages = max_pages

    texts = []
    processed_pages = 0
    for page_num in range(1, total_pages + 1):
        try:
            images = pdf2image.convert_from_path(
                tmp_pdf_path,
                first_page=page_num,
                last_page=page_num,
                dpi=max(72, OCR_RENDER_DPI),
                grayscale=True,
                fmt="jpeg",
                thread_count=1,
            )
        except Exception:
            break
        if not images:
            break
        img = images[0]
        text = pytesseract.image_to_string(img)
        if text:
            texts.append(text)
        try:
            img.close()
        except Exception:
            pass
        del images
        gc.collect()
        processed_pages += 1
    elapsed = time.perf_counter() - t0
    logger.info("[extract] OCR duration=%.2fs pages=%d", elapsed, processed_pages)
    if tmp_pdf_path:
        try:
            os.unlink(tmp_pdf_path)
        except Exception:
            pass
    return "\n\n".join(texts) if texts else ""


def extract_text_from_pdf_with_ocr(
    file: BinaryIO,
    force_ocr: bool = False,
    ocr_pages: int = DEFAULT_OCR_PAGES,
) -> tuple[str, str]:
    """
    Extract text from PDF. Run OCR only when needed; limit to first ocr_pages.
    Returns (combined_text, source_string).
    """
    t0 = time.perf_counter()
    pdf_text = extract_text_from_pdf(file)
    elapsed = time.perf_counter() - t0
    logger.info("[extract] PDF text extraction duration=%.2fs len=%d", elapsed, len(pdf_text))

    # Early exit: enough text and high quality -> no OCR
    if not force_ocr and len(pdf_text.strip()) >= OCR_EARLY_EXIT_MIN_CHARS and _text_looks_high_quality(pdf_text):
        return (pdf_text or "").strip(), "pdf_text"

    run_ocr = force_ocr or len(pdf_text.strip()) < OCR_EARLY_EXIT_MIN_CHARS or not _text_looks_high_quality(pdf_text)
    if not run_ocr:
        return (pdf_text or "").strip(), "pdf_text"

    file.seek(0)
    ocr_text = _run_ocr_on_pdf(file, max_pages=ocr_pages)
    if not pdf_text.strip():
        return (ocr_text or "").strip(), "ocr"
    combined = "=== PDF TEXT ===\n\n" + pdf_text.strip() + "\n\n=== OCR TEXT ===\n\n" + (ocr_text or "").strip()
    return combined.strip(), "pdf_text+ocr"


def extract_text_from_docx(file: BinaryIO) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")
    try:
        raw_bytes = file.read()
        file.seek(0)
        doc = Document(file)
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Proposal/LOI files often place key economics in tables. Include them in text payload.
        for table in doc.tables:
            for row in table.rows:
                cells = [" ".join((cell.text or "").split()) for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))
        # Build "accepted changes" text directly from WordprocessingML.
        # Keep inserted/current text, ignore deleted text so redlined LOIs parse the final terms.
        try:
            with zipfile.ZipFile(BytesIO(raw_bytes)) as zf:
                xml = zf.read("word/document.xml")
            root = ET.fromstring(xml)
            w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

            def collect_text_nodes(node: ET.Element, in_deleted: bool = False) -> list[str]:
                current_deleted = in_deleted or node.tag == f"{w_ns}del"
                out: list[str] = []
                if node.tag == f"{w_ns}t" and not current_deleted:
                    txt = node.text or ""
                    if txt:
                        out.append(txt)
                elif node.tag in (f"{w_ns}tab", f"{w_ns}br", f"{w_ns}cr") and not current_deleted:
                    # Preserve structural separators from WordprocessingML so labels/values
                    # don't collapse into merged tokens like "ImprovementsLandlord".
                    out.append(" ")
                for child in list(node):
                    out.extend(collect_text_nodes(child, current_deleted))
                return out

            xml_lines: list[str] = []
            for p_node in root.iter(f"{w_ns}p"):
                tokens = collect_text_nodes(p_node, False)
                if not tokens:
                    continue
                line = "".join(tokens)
                line = re.sub(r"\s+", " ", line).strip()
                if line:
                    xml_lines.append(line)

            if xml_lines:
                parts.extend(xml_lines)
        except Exception:
            pass

        def _extract_docx_image_ocr_chunks(docx_bytes: bytes) -> list[str]:
            try:
                from PIL import Image
                import pytesseract
            except Exception:
                return []

            image_chunks: list[str] = []
            try:
                with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
                    media_files = [
                        name
                        for name in zf.namelist()
                        if name.startswith("word/media/")
                        and name.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"))
                    ]
                    if not media_files:
                        return []

                    candidates: list[tuple[int, str, bytes]] = []
                    for name in media_files:
                        try:
                            data = zf.read(name)
                            if not data:
                                continue
                            with Image.open(BytesIO(data)) as img:
                                width, height = img.size
                            area = int(width) * int(height)
                            if area < DOCX_OCR_MIN_IMAGE_AREA:
                                continue
                            candidates.append((area, name, data))
                        except Exception:
                            continue
            except Exception:
                return []

            if not candidates:
                return []

            candidates.sort(key=lambda row: row[0], reverse=True)
            bytes_budget = DOCX_OCR_MAX_TOTAL_IMAGE_BYTES
            taken = 0
            for _area, name, data in candidates:
                if taken >= DOCX_OCR_MAX_IMAGES:
                    break
                if bytes_budget <= 0:
                    break
                if len(data) > bytes_budget:
                    continue
                try:
                    from PIL import Image  # keep import local for optional dependency behavior
                    import pytesseract

                    with Image.open(BytesIO(data)) as img:
                        img_l = img.convert("L")
                        width, height = img_l.size
                        if max(width, height) < 1200:
                            scale = 2
                            img_l = img_l.resize((max(1, width * scale), max(1, height * scale)))
                        ocr_raw = pytesseract.image_to_string(img_l, config="--oem 1 --psm 11")
                    ocr_text = _clean_word_text(ocr_raw)
                    if len(ocr_text) < 80:
                        continue
                    if sum(ch.isdigit() for ch in ocr_text) < 8:
                        continue
                    image_chunks.append(f"[DOCX IMAGE OCR {Path(name).name}]\n{ocr_text}")
                    bytes_budget -= len(data)
                    taken += 1
                except Exception:
                    continue
            return image_chunks

        # OCR embedded DOCX images only when key economics look incomplete.
        current_text = "\n\n".join(parts)
        current_low = current_text.lower()
        has_base_rent_amount = bool(
            re.search(
                r"(?is)\b(?:base\s+(?:annual\s+net\s+)?rental?\s+rate|base\s+rent)\b[^\n]{0,140}\$\s*[\d,]+(?:\.\d+)?",
                current_text,
            )
        )
        has_parking_language = "parking" in current_low
        has_structured_parking_data = bool(
            re.search(
                r"(?i)\b(?:parking\s+ratio|#\s*reserved\s+(?:paid\s+)?spaces?|#\s*unreserved\s+(?:paid\s+)?spaces?|total\s*#?\s*paid\s+spaces)\b"
                r"|(?:\d+(?:\.\d+)?)\s*(?:/|per)\s*1,?000\s*(?:rsf|sf)\b",
                current_text,
            )
        )
        has_explicit_escalation = bool(
            re.search(
                r"(?i)\b(?:annual\s+base\s+rent\s+escalation|annual\s+rental\s+increases?|base\s+rent\b[^\n]{0,140}\d+(?:\.\d+)?%\s*(?:increase|escalat))\b",
                current_text,
            )
        )
        should_ocr_images = bool(
            re.search(r"(?i)\b(as\s+shown\s+on\s+the\s+attached\s+schedule|attached\s+schedule|rent\s+schedule)\b", current_text)
        ) or not has_base_rent_amount or (has_parking_language and not has_structured_parking_data) or not has_explicit_escalation
        if should_ocr_images:
            parts.extend(_extract_docx_image_ocr_chunks(raw_bytes))

        return "\n\n".join(parts) if parts else ""
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}") from e


def _clean_word_text(text: str) -> str:
    text = (text or "").replace("\ufeff", "").replace("\x00", " ").replace("\r", "\n")
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue
        lines.append(line)
    return "\n\n".join(lines)


def _decode_command_output(raw: bytes) -> str:
    decoded = raw.decode("utf-8", errors="ignore").strip()
    if not decoded:
        decoded = raw.decode("latin-1", errors="ignore").strip()
    return _clean_word_text(decoded)


def _extract_printable_text_fallback(raw_bytes: bytes) -> str:
    decoded = raw_bytes.decode("latin-1", errors="ignore")
    decoded = decoded.replace("\x00", " ").replace("\r", "\n")
    lines: list[str] = []
    for raw_line in decoded.splitlines():
        line = re.sub(r"[^A-Za-z0-9$%/.,:;()&@#'\"!?+\- ]+", " ", raw_line)
        line = re.sub(r"\s+", " ", line).strip()
        if len(line) < 20:
            continue
        if not re.search(r"[A-Za-z]{3,}", line):
            continue
        lines.append(line)
    if not lines:
        return ""
    deduped: list[str] = []
    seen: set[str] = set()
    for line in lines:
        if line in seen:
            continue
        seen.add(line)
        deduped.append(line)
        if len(deduped) >= 400:
            break
    return _clean_word_text("\n".join(deduped))


def _run_text_command(command: list[str]) -> str:
    try:
        proc = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=DOC_TEXT_COMMAND_TIMEOUT_SECONDS,
            check=False,
        )
    except Exception:
        return ""
    if proc.returncode != 0 and not proc.stdout:
        return ""
    return _decode_command_output(proc.stdout)


def _extract_text_with_soffice(source_path: str) -> str:
    if not shutil.which("soffice"):
        return ""
    with tempfile.TemporaryDirectory(prefix="word-soffice-") as tmpdir:
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "txt:Text",
            "--outdir",
            tmpdir,
            source_path,
        ]
        try:
            subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=DOC_TEXT_COMMAND_TIMEOUT_SECONDS,
                check=False,
            )
        except Exception:
            return ""
        txt_path = Path(tmpdir) / f"{Path(source_path).stem}.txt"
        if not txt_path.exists():
            return ""
        try:
            return _clean_word_text(txt_path.read_text(encoding="utf-8", errors="ignore"))
        except Exception:
            return ""


def extract_text_from_doc(file: BinaryIO, filename: str = "") -> str:
    """
    Extract text from legacy .doc files using command-line converters with fallbacks.
    """
    try:
        raw_bytes = file.read()
        file.seek(0)
    except Exception as e:
        raise ValueError(f"Failed to read DOC: {e}") from e
    if not raw_bytes:
        return ""

    # Some uploads are mislabeled .doc but are actually OOXML containers.
    if raw_bytes.startswith(b"PK"):
        try:
            return extract_text_from_docx(BytesIO(raw_bytes))
        except Exception:
            pass

    suffix = ".doc"
    lower_name = (filename or "").strip().lower()
    if lower_name.endswith(".rtf"):
        suffix = ".rtf"
    elif lower_name.endswith(".txt"):
        suffix = ".txt"

    with tempfile.NamedTemporaryFile(prefix="word-upload-", suffix=suffix, delete=False) as tmp:
        tmp.write(raw_bytes)
        tmp_path = tmp.name

    try:
        for command in (
            ["antiword", tmp_path],
            ["catdoc", tmp_path],
            ["textutil", "-convert", "txt", "-stdout", tmp_path],
        ):
            if not shutil.which(command[0]):
                continue
            text = _run_text_command(command)
            if len(text) >= MIN_DOC_TEXT_CHARS:
                return text

        soffice_text = _extract_text_with_soffice(tmp_path)
        if len(soffice_text) >= MIN_DOC_TEXT_CHARS:
            return soffice_text

        if suffix == ".txt":
            plain = _clean_word_text(raw_bytes.decode("utf-8", errors="ignore"))
            if plain:
                return plain

        fallback = _extract_printable_text_fallback(raw_bytes)
        if fallback:
            return fallback
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
    raise ValueError(
        "Failed to read DOC. Install at least one converter (antiword, catdoc, or LibreOffice)."
    )


def extract_text_from_word(file: BinaryIO, filename: str = "") -> tuple[str, str]:
    """
    Route to DOCX or DOC extraction based on extension + file signature.
    Returns (text, source) where source is "docx" or "doc".
    """
    raw_bytes = file.read()
    file.seek(0)
    lower_name = (filename or "").strip().lower()
    is_docx = lower_name.endswith(".docx") or raw_bytes.startswith(b"PK")
    if is_docx:
        return extract_text_from_docx(BytesIO(raw_bytes)), "docx"
    return extract_text_from_doc(BytesIO(raw_bytes), filename=filename), "doc"


def _build_snippet_pack(text: str) -> tuple[str, bool]:
    """
    Build snippet pack: lines within +/- SNIPPET_CONTEXT_LINES of keyword matches.
    Dedupe and cap at SNIPPET_PACK_MAX_CHARS. Returns (pack_text, True if used).
    """
    t0 = time.perf_counter()
    lines = text.splitlines()
    keywords_lower = [k.lower() for k in FOCUS_KEYWORDS]
    match_indices = set()
    for i, line in enumerate(lines):
        ll = line.lower()
        if any(kw in ll for kw in keywords_lower):
            for j in range(
                max(0, i - SNIPPET_CONTEXT_LINES),
                min(len(lines), i + SNIPPET_CONTEXT_LINES + 1),
            ):
                match_indices.add(j)
    if not match_indices:
        elapsed = time.perf_counter() - t0
        logger.info("[extract] snippet build duration=%.2fs (no matches, fallback)", elapsed)
        return text, False
    # Build paragraphs (contiguous line blocks), dedupe by content
    seen = set()
    chunks = []
    for idx in sorted(match_indices):
        block = lines[idx].strip()
        if not block:
            continue
        if block in seen:
            continue
        seen.add(block)
        chunks.append(block)
    pack = "\n\n".join(chunks)
    if len(pack) > SNIPPET_PACK_MAX_CHARS:
        pack = pack[:SNIPPET_PACK_MAX_CHARS] + "\n\n...[snippet pack truncated]..."
    elapsed = time.perf_counter() - t0
    logger.info("[extract] snippet build duration=%.2fs chars=%d", elapsed, len(pack))
    if len(pack) < SNIPPET_MIN_CHARS:
        return text, False
    return pack, True


def _truncate_fallback(text: str) -> tuple[str, list[str]]:
    """First 20k + last 20k with separator. Returns (text, extra_warnings)."""
    warnings = []
    if len(text) <= FALLBACK_HEAD_CHARS + FALLBACK_TAIL_CHARS:
        return text, warnings
    head = text[:FALLBACK_HEAD_CHARS]
    tail = text[-FALLBACK_TAIL_CHARS:] if len(text) > FALLBACK_TAIL_CHARS else ""
    truncated = head + "\n\n...[truncated]...\n\n" + tail
    warnings.append("Input text truncated for extraction.")
    return truncated, warnings


def _prepare_text_for_llm(text: str) -> tuple[str, list[str]]:
    """Snippet pack first; if too small, fallback truncation. Returns (text_for_llm, extra_warnings)."""
    snippet, used = _build_snippet_pack(text)
    if used:
        return _truncate_fallback(snippet)
    return _truncate_fallback(text)


# ---- Regex pre-extraction ----
_SF_UNIT_PATTERN = r"(?:r\.?s\.?f\.?|s\.?f\.?|rsf|sf|psf|sq\.?\s*ft\.?|square\s*(?:feet?|foot))"
_RE_RSF = re.compile(
    r"\b(?:r\.?s\.?f\.?|rsf|rentable\s+(?:square\s+(?:feet?|foot)|s\.?f\.?|sf)|square\s*(?:feet?|foot)|sq\.?\s*ft\.?)\b[ \t:]*[\d,]+\.?\d*|"
    rf"[\d,]+\.?\d*[ \t]*(?:{_SF_UNIT_PATTERN}|rentable\s+(?:square\s+(?:feet?|foot)|s\.?f\.?|sf))\b",
    re.I,
)
_RE_NUM = re.compile(r"[\d,]+\.?\d*")
_RE_ISO_DATE = re.compile(r"\b(20\d{2})[-/](0?[1-9]|1[0-2])[-/](0?[1-9]|[12]\d|3[01])\b")
_RE_US_DATE = re.compile(r"\b(0?[1-9]|1[0-2])[/-](0?[1-9]|[12]\d|3[01])[/-](20\d{2})\b")
_RE_MONTH_DATE = re.compile(
    r"(?i)\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
    r"aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*,?\s*(\d{1,2})(?:st|nd|rd|th)?,?\s+(20\d{2})\b"
)
_RE_COMM_LABEL = re.compile(
    r"(?i)\b(?:estimated\s+commencement(?:\s+date)?|commencement(?:\s+date)?|commencing|lease\s+commencement)\b[^A-Za-z0-9]{0,20}([^\n]{0,120})"
)
_RE_EXP_LABEL = re.compile(
    r"(?i)\b(?:estimated\s+(?:termination|expiration)\s+date|termination(?:\s+date)?|expiration(?:\s+date)?|expires?|expiring|lease\s+expiration)\b[^A-Za-z0-9]{0,20}([^\n]{0,120})"
)
_RE_BUILDING = re.compile(r"(?i)\b(?:building|property)\s*(?:name)?\s*[:#-]\s*([^\n,;]{3,80})")
_RE_SUITE = re.compile(r"(?i)\b(?:suite|ste\.?|unit|space|premises)\b\s*[:#-]?\s*([A-Za-z0-9][A-Za-z0-9\- ]{0,30})")
_RE_TI = re.compile(
    r"(?i)\b(?:ti\s+allowance|tenant\s+allowance|tenant\s+improvement(?:s)?\s+allowance|improvement\s+allowance)\b[^\n$]{0,260}\$?\s*[\d,]+(?:\.\d+)?",
    re.I,
)
_RE_TIA_KEYWORD = re.compile(
    r"(?i)\b(?:"
    r"tia|"
    r"ti\s+allowance|"
    r"tenant\s+allowance|"
    r"subtenant\s+allowance|"
    r"tenant\s+improvement(?:s)?(?:\s+allowance)?|"
    r"subtenant\s+improvement(?:s)?(?:\s+allowance)?|"
    r"improvement\s+allowance|"
    r"allowance\s*&\s*improvements?"
    r")\b"
)
_RE_TI_ALLOWANCE_TOTAL = re.compile(
    r"(?i)\b(?:ti(?:a)?|ti\s+allowance|tenant\s+allowance|tenant\s+improvement(?:s)?\s+allowance|improvement\s+allowance)\b"
    r"[^\n$]{0,260}\$\s*([\d,]+(?:\.\d+)?)"
)
_RE_PSF_AMOUNT = re.compile(
    rf"(?i)\$?\s*([\d,]+(?:\.\d+)?)\s*(?:/|per)?\s*(?:rentable\s+)?{_SF_UNIT_PATTERN}\b"
)
_RE_FREE_RENT = re.compile(
    r"(?i)\b(?:free\s+rent|rent\s+abatement|abatement|abated)\b[^\n]{0,50}?\(?(\d{1,2})\)?\s*(?:months?)?",
    re.I,
)
_RE_BASE_OPEX = re.compile(
    r"(?i)\b(?:base\s+year\s+)?(?:opex|ope\b|operating\s+expenses?|cam)\b[^\n]{0,220}",
    re.I,
)
_RE_BASE_RENT = re.compile(
    r"\b(?:lease\s+rate|base\s+rent|basic\s+rent|rental\s+rate|annual\s+rent)\b[^\n$]{0,40}\$?\s*([\d,]+\.?\d*)",
    re.I,
)
_RE_BASE_RENT_FLEX = re.compile(
    rf"(?is)\b(?:initial\s+)?(?:lease\s+rate|base\s+rent|basic\s+rent|rental\s+rate|annual\s+rent)\b.{{0,240}}?\$\s*([\d,]+(?:\.\d{{1,4}})?)\s*(?:/|per)?\s*(?:rentable\s+)?{_SF_UNIT_PATTERN}\b"
)
_RE_YEAR_RATE_INLINE = re.compile(
    r"(?i)\b(?:lease\s*)?years?\s*(\d{1,2})(?:\s*(?:-|to|through|thru|–|—)\s*(\d{1,2}))?"
    r"\b[^\n$]{0,120}\$?\s*([\d,]+(?:\.\d{1,4})?)"
)
_RE_YEAR_RATE_INLINE_MONTHLY_PSF = re.compile(
    rf"(?i)\b(?:lease\s*)?years?\s*(\d{{1,2}})(?:\s*(?:-|to|through|thru|–|—)\s*(\d{{1,2}}))?"
    rf"\b[^\n]{{0,100}}?\$\s*[\d,]+(?:\.\d{{1,4}})?\s*(?:per|/)\s*(?:mo(?:nth)?s?|mos?\.?)\b"
    rf"[^\n]{{0,60}}?\(\s*\$?\s*([\d,]+(?:\.\d{{1,4}})?)\s*(?:/|per)\s*(?:rentable\s+)?{_SF_UNIT_PATTERN}\s*\)"
)
_RE_YEAR_TABLE_HEADER = re.compile(r"(?i)\b(?:lease\s*)?years?\b.*\b(?:rent|rate|psf|/sf|per\s+sf)\b")
_RE_YEAR_TABLE_ROW = re.compile(
    r"^\s*(\d{1,2})(?:\s*(?:-|to|through|thru|–|—)\s*(\d{1,2}))?\s+\$?\s*([\d,]+(?:\.\d{1,4})?)\b"
)
_RE_YEAR_LABEL_ONLY = re.compile(
    r"(?i)^\s*(?:lease\s*)?years?\s*(\d{1,2})(?:\s*(?:-|to|through|thru|–|—)\s*(\d{1,2}))?\s*$"
)
_RE_SIMPLE_RATE_LINE = re.compile(r"^\s*\$?\s*([\d,]+(?:\.\d{1,4})?)\s*$")
_RE_RENT_TABLE_HINT = re.compile(
    r"(?i)\b(base\s+rent|lease\s+year|annual\s+rate|rate\s*\(psf\)|monthly\s+base\s+rent)\b"
)
_RE_NON_RENT_YEAR_CONTEXT = re.compile(
    r"(?i)\b(parking|space(?:s)?|per\s+space|free\s+rent|abatement|opex|operating\s+expense|cam|tax(?:es)?)\b"
)
_RE_MONTH_RANGE_RATE = re.compile(
    r"(?i)\$\s*([\d,]+(?:\.\d{1,4})?)\s*(?:per|/)\s*(?:rsf|sf|psf)\b[^\n]{0,120}\bfor\s+months?\s*(\d{1,3})\s*(?:-|to|through|thru|–|—)\s*(\d{1,3})\b"
)
_RE_ANNUAL_ESCALATION = re.compile(
    r"(?i)\b(\d+(?:\.\d+)?)%\s*"
    r"(?:(?:annual(?:ly)?\s+)?(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
    r"|annual(?:ly)?\s+(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b)"
    r"[^\n]{0,120}\b(?:starting|beginning|commencing)\s*(?:in\s*)?month\s*(\d{1,3})\b"
)
_RE_ANNUAL_ESCALATION_GENERIC = re.compile(
    r"(?i)\b(\d+(?:\.\d+)?)%\s+"
    r"(?:(?:annual(?:ly)?\s+(?:(?:base\s+)?rent(?:al)?(?:\s+rate)?\s+)?)?(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
    r"|(?:increase(?:s)?|escalat(?:ion|ions)|escalator)\b[^\n]{0,60}\bannual(?:ly| anniversary)?\b)"
)
_RE_ANNUAL_ESCALATION_LABEL_VALUE = re.compile(
    r"(?i)\bannual(?:\s+(?:base\s+)?rent(?:al)?(?:\s+rate)?)?\s+"
    r"(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
    r"(?:\s*[:|=]\s*|[^\n]{0,140}?)(\d+(?:\.\d+)?)\s*%"
)
_RE_BASE_RENT_INCREASE_PERCENT_AFTER = re.compile(
    r"(?i)\b(?:base\s+rent(?:al)?(?:\s+rate)?|rent(?:al)?\s+rate)\b[^\n]{0,120}?"
    r"\b(?:increase(?:s|d)?|escalat(?:e|ed|ion|ions)|escalator)\b[^\n%]{0,80}?"
    r"(\d+(?:\.\d+)?)\s*%\s*(?:per\s+year|annually|annual(?:ly)?)?"
)
_RE_BASE_RENT_PERCENT_INCREASE_ANNUAL_LATER = re.compile(
    r"(?i)\b(?:base\s+rent(?:al)?(?:\s+rate)?|rent(?:al)?\s+rate)\b[^\n]{0,180}?"
    r"(\d+(?:\.\d+)?)\s*%\s*(?:increase(?:s|d)?|escalat(?:e|ed|ion|ions)|escalator)\b"
    r"[^\n]{0,80}\bannual(?:ly| anniversary)?\b"
)
_RE_ANNUAL_ESCALATION_PERCENT_IN_PARENS = re.compile(
    r"(?i)\(\s*(\d+(?:\.\d+)?)%\s*\)\s*"
    r"(?:(?:annual(?:ly)?\s+)?(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
    r"|annual(?:ly)?\s+(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b)"
)
_RE_ANNUAL_ESCALATION_START_IN_PARENS = re.compile(
    r"(?i)\(\s*(\d+(?:\.\d+)?)%\s*\)\s*"
    r"(?:(?:annual(?:ly)?\s+)?(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
    r"|annual(?:ly)?\s+(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b)"
    r"[^\n]{0,120}\b(?:starting|beginning|commencing)\s*(?:in\s*)?month\s*(\d{1,3})\b"
)


_WORD_UNITS = {
    "zero": 0,
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
    "eleven": 11,
    "twelve": 12,
    "thirteen": 13,
    "fourteen": 14,
    "fifteen": 15,
    "sixteen": 16,
    "seventeen": 17,
    "eighteen": 18,
    "nineteen": 19,
}
_WORD_TENS = {
    "twenty": 20,
    "thirty": 30,
    "forty": 40,
    "fifty": 50,
    "sixty": 60,
    "seventy": 70,
    "eighty": 80,
    "ninety": 90,
}


def _word_token_to_int(token: str | None) -> int | None:
    if not token:
        return None
    cleaned = re.sub(r"[^a-z\- ]", "", token.lower()).strip()
    if not cleaned:
        return None
    parts = [p for p in re.split(r"[-\s]+", cleaned) if p]
    if not parts:
        return None
    if len(parts) == 1:
        if parts[0] in _WORD_UNITS:
            return _WORD_UNITS[parts[0]]
        if parts[0] in _WORD_TENS:
            return _WORD_TENS[parts[0]]
        return None
    if len(parts) == 2 and parts[0] in _WORD_TENS and parts[1] in _WORD_UNITS:
        return _WORD_TENS[parts[0]] + _WORD_UNITS[parts[1]]
    return None


def _percent_token_to_float(token: str | None) -> float | None:
    pct = _coerce_float_token(token, None)
    if pct is not None:
        return float(pct)
    word_val = _word_token_to_int(token)
    if word_val is not None:
        return float(word_val)
    cleaned = re.sub(r"[^a-z\- ]", " ", str(token or "").lower()).strip()
    if not cleaned:
        return None
    parts = [p for p in re.split(r"[-\s]+", cleaned) if p]
    for width in (2, 1):
        if len(parts) < width:
            continue
        for end in range(len(parts), width - 1, -1):
            candidate = " ".join(parts[end - width:end])
            parsed = _word_token_to_int(candidate)
            if parsed is not None:
                return float(parsed)
    return None


def _has_ti_context_token(text: str) -> bool:
    segment = str(text or "")
    if not segment:
        return False
    if _RE_TIA_KEYWORD.search(segment):
        return True
    compact = re.sub(r"[^a-z0-9]+", "", segment.lower())
    return (
        "tenantimprovementallowance" in compact
        or "tenantimprovementsallowance" in compact
        or "tenantimprovement" in compact
        or "tenantimprovements" in compact
    )


def _extract_annual_rent_escalation_pct(text: str) -> float | None:
    if not text:
        return None
    candidates: list[tuple[int, float, int]] = []
    patterns = (
        (_RE_ANNUAL_ESCALATION_PERCENT_IN_PARENS, 1, 5),
        (
            re.compile(
                r"(?i)\b([a-z]+(?:[-\s]+[a-z]+)?)\s+percent\s+"
                r"(?:(?:annual(?:ly)?\s+)?(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
                r"|annual(?:ly)?\s+(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b)"
            ),
            1,
            3,
        ),
        (_RE_ANNUAL_ESCALATION_GENERIC, 1, 4),
        (
            re.compile(
                r"(?i)\bannual(?:\s+(?:base\s+)?rent(?:al)?(?:\s+rate)?)?\s+"
                r"(?:escalat(?:ion|ions)|increase(?:s)?|escalator)\b"
                r"(?:\s*[:|=]\s*|[^\n]{0,140}?)([a-z]+(?:[-\s]+[a-z]+)?)\s+percent\b"
            ),
            1,
            3,
        ),
        (_RE_ANNUAL_ESCALATION_LABEL_VALUE, 1, 4),
        (_RE_BASE_RENT_INCREASE_PERCENT_AFTER, 1, 5),
        (_RE_BASE_RENT_PERCENT_INCREASE_ANNUAL_LATER, 1, 6),
    )
    for pattern, group_index, base_score in patterns:
        for match in pattern.finditer(text):
            pct = _percent_token_to_float(match.group(group_index))
            if pct is None or pct <= 0 or pct > 25:
                continue
            segment = (match.group(0) or "").lower()
            score = base_score
            if any(tok in segment for tok in ("base rent", "rental rate", "rent escalation", "rent increase")):
                score += 4
            if any(tok in segment for tok in ("opex", "operating expense", "cam")):
                score -= 6
            candidates.append((score, float(pct), match.start()))
    if not candidates:
        return None
    candidates.sort(key=lambda row: (-row[0], row[2]))
    best_score, best_pct, _ = candidates[0]
    if best_score <= 0:
        return None
    return best_pct


def _extract_annual_rent_escalation_start_month(text: str) -> int | None:
    if not text:
        return None
    candidates: list[tuple[int, int, int]] = []
    patterns = (
        (_RE_ANNUAL_ESCALATION_START_IN_PARENS, 1, 2, 7),
        (_RE_ANNUAL_ESCALATION, 1, 2, 6),
    )
    for pattern, pct_group, month_group, base_score in patterns:
        for match in pattern.finditer(text):
            pct = _coerce_float_token(match.group(pct_group), None)
            month = _coerce_int_token(match.group(month_group), None)
            if pct is None or pct <= 0 or pct > 25:
                continue
            if month is None or month <= 0 or month > 600:
                continue
            score = base_score
            seg = (match.group(0) or "").lower()
            if "base rent" in seg or "rental rate" in seg or "lease rate" in seg:
                score += 2
            candidates.append((score, int(month), match.start()))
    if not candidates:
        return None
    candidates.sort(key=lambda row: (-row[0], row[2], row[1]))
    return int(candidates[0][1])


def _parse_text_date_token(s: str) -> str | None:
    token = (s or "").strip()
    if not token:
        return None
    token = re.sub(r"(?i)\b(\d{1,2})(st|nd|rd|th)\b", r"\1", token)
    token = re.sub(
        r"(?i)\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|"
        r"sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s*,\s*(\d{1,2})",
        r"\1 \2",
        token,
    )
    token = re.sub(r"\s+", " ", token).strip()
    m = _RE_ISO_DATE.search(token)
    if m:
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return f"{y:04d}-{mo:02d}-{d:02d}"
    m = _RE_US_DATE.search(token)
    if m:
        mo, d, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return f"{y:04d}-{mo:02d}-{d:02d}"
    m = _RE_MONTH_DATE.search(token)
    if m:
        mon = m.group(1).lower()[:3]
        months = {
            "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
            "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
        }
        mo = months.get(mon)
        if mo:
            d = int(m.group(2))
            y = int(m.group(3))
            return f"{y:04d}-{mo:02d}-{d:02d}"
    return None


def _coerce_int_token(value: Any, default: int | None = None) -> int | None:
    if value is None or isinstance(value, bool):
        return default
    if isinstance(value, (int, float)):
        if value != value:
            return default
        return int(value)
    if isinstance(value, str):
        token = value.strip().replace(",", "")
        if not token:
            return default
        m = re.search(r"-?\d+", token)
        if m:
            try:
                return int(m.group(0))
            except ValueError:
                return default
    return default


def _coerce_float_token(value: Any, default: float | None = None) -> float | None:
    if value is None or isinstance(value, bool):
        return default
    if isinstance(value, (int, float)):
        if value != value:
            return default
        return float(value)
    if isinstance(value, str):
        token = value.strip().replace(",", "")
        if not token:
            return default
        m = re.search(r"-?\d+(?:\.\d+)?", token)
        if m:
            try:
                return float(m.group(0))
            except ValueError:
                return default
    return default


def _parking_count_from_ratio_per_1000(parking_ratio: Any, rsf: Any) -> int:
    ratio_val = max(0.0, _coerce_float_token(parking_ratio, 0.0) or 0.0)
    rsf_val = max(0.0, _coerce_float_token(rsf, 0.0) or 0.0)
    if ratio_val <= 0 or rsf_val <= 0:
        return 0
    return max(0, int(round((ratio_val * rsf_val) / 1000.0)))


def _parking_ratio_per_1000_from_count(parking_count: Any, rsf: Any) -> float:
    count_val = max(0, _coerce_int_token(parking_count, 0) or 0)
    rsf_val = max(0.0, _coerce_float_token(rsf, 0.0) or 0.0)
    if count_val <= 0 or rsf_val <= 0:
        return 0.0
    return float((float(count_val) * 1000.0) / rsf_val)


def _reconcile_parking_ratio_and_count(
    *,
    parking_ratio: Any,
    parking_count: Any,
    rsf: Any,
    prefer: str = "ratio",
) -> tuple[float, int]:
    ratio_val = max(0.0, _coerce_float_token(parking_ratio, 0.0) or 0.0)
    count_val = max(0, _coerce_int_token(parking_count, 0) or 0)
    rsf_val = max(0.0, _coerce_float_token(rsf, 0.0) or 0.0)
    if rsf_val <= 0:
        return float(ratio_val), int(count_val)

    derived_count = _parking_count_from_ratio_per_1000(ratio_val, rsf_val) if ratio_val > 0 else 0
    derived_ratio = _parking_ratio_per_1000_from_count(count_val, rsf_val) if count_val > 0 else 0.0
    prefer_source = "count" if str(prefer or "").strip().lower() == "count" else "ratio"

    if ratio_val > 0 and count_val <= 0 and derived_count > 0:
        count_val = derived_count
    elif count_val > 0 and ratio_val <= 0 and derived_ratio > 0:
        ratio_val = derived_ratio
    elif ratio_val > 0 and count_val > 0 and derived_count > 0:
        if abs(derived_count - count_val) > 1:
            if prefer_source == "count" and derived_ratio > 0:
                ratio_val = derived_ratio
            elif prefer_source == "ratio":
                count_val = derived_count

    return float(ratio_val), int(count_val)


def _term_month_count(commencement: date, expiration: date) -> int:
    """
    Return lease term as month count.
    Example: 2026-01-01 to 2031-01-31 => 60 (months 0..59).
    """
    months = (expiration.year - commencement.year) * 12 + (expiration.month - commencement.month)
    if expiration.day < commencement.day:
        months -= 1
    return max(1, months)


def _expiration_from_term_months(commencement: date, term_months: int) -> date:
    """Compute lease expiration from commencement + term months (inclusive month accounting)."""
    tm = max(1, int(term_months))
    total = (commencement.month - 1) + tm
    year = commencement.year + (total // 12)
    month = (total % 12) + 1
    day = min(commencement.day, monthrange(year, month)[1])
    anniv = date(year, month, day)
    return anniv - timedelta(days=1)


def _looks_like_abatement_distribution_window(text: str) -> bool:
    low = str(text or "").lower()
    if not low:
        return False
    if not any(tok in low for tok in ("abatement", "abated", "abate", "free rent")):
        return False
    has_distribution_language = any(
        tok in low
        for tok in (
            "spread",
            "installment",
            "installments",
            "throughout",
            "amortiz",
            "pro rata",
            "prorata",
            "over the first",
            "over first",
        )
    )
    if not has_distribution_language:
        return False
    return bool(
        re.search(
            r"(?i)\bfirst\s+(?:[a-z\-]+\s*)?\(?\d{1,3}\)?\s+months?\b|\bmonths?\s+1\s*(?:-|to|through|thru|–|—)\s*\d{1,3}\b",
            low,
        )
    )


def _extract_term_months_from_text(text: str) -> int | None:
    if not text:
        return None
    disqualifying_tokens = (
        "extension",
        "renewal",
        "option",
        "additional",
        "holdover",
        "prior written notice",
        "month-to-month",
        "prior to",
        "no earlier than",
        "no later than",
    )

    def _match_context(match: re.Match[str]) -> str:
        start = match.start()
        end = match.end()
        line_start = text.rfind("\n", 0, start) + 1
        line_end = text.find("\n", end)
        if line_end == -1:
            line_end = len(text)
        return text[line_start:line_end].lower()

    def _is_disqualified(match: re.Match[str]) -> bool:
        context = _match_context(match)
        if any(token in context for token in disqualifying_tokens):
            return True
        if _looks_like_abatement_distribution_window(context):
            return True
        return False

    def _looks_like_non_term_month_context(snippet: str) -> bool:
        low = str(snippet or "").lower()
        if not low:
            return False
        if _looks_like_abatement_distribution_window(low):
            return True
        if re.search(r"(?i)\bfirst\s+\(?\d{1,3}\)?\s+months?\s+of\s+the\s+term\b", low):
            return True
        if re.search(r"(?i)\b\d{1,3}\s+months?\s+prior\b", low):
            return True
        if re.search(r"(?i)\b(?:starting|beginning|commencing)\s+month\s+\d{1,3}\b", low):
            return True
        if (
            re.search(r"(?i)\bmonth\s+\d{1,3}\b", low)
            and any(tok in low for tok in ("escalat", "increase", "base rent", "rental rate"))
        ):
            return True
        if any(
            tok in low
            for tok in (
                "parking",
                "access card",
                "parking pass",
                "must take and pay",
                "event parking",
            )
        ) and "lease term" not in low and "initial term" not in low:
            return True
        return False

    candidates: list[tuple[int, int, int]] = []

    # Handle split-line term blocks:
    #   RENEWAL TERM:
    #   Ninety(90) months
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    term_label_pat = re.compile(
        r"(?i)\b(?:primary\s+lease\s+term|lease\s+term|sublease\s+term|renewal\s+term|initial\s+term)\b|"
        r"\bterm\s+and\s+free\s+rent\b|"
        r"\bterm\s*[:\-]|"
        r"\bterm\s+(?:shall\s+be|is|of)\b"
    )
    for idx, line in enumerate(lines[:1400]):
        low_line = line.lower()
        if not term_label_pat.search(line):
            continue
        if any(tok in low_line for tok in ("option to extend", "option term", "renewal option", "holdover")):
            continue
        if _looks_like_non_term_month_context(low_line):
            continue
        for lookahead in range(0, 4):
            if idx + lookahead >= len(lines):
                break
            probe = lines[idx + lookahead]
            probe_low = probe.lower()
            if any(tok in probe_low for tok in disqualifying_tokens):
                continue
            if _looks_like_non_term_month_context(probe_low):
                continue
            month_matches = re.findall(
                r"(?i)\b(?:[a-z\-]+\s*)?\(?(\d{1,3})\)?\s*(?:calendar\s+)?months?\b",
                probe,
            )
            if month_matches:
                parsed_months = [_coerce_int_token(token, 0) for token in month_matches]
                parsed_months = [m for m in parsed_months if 1 <= m <= 600]
                if not parsed_months:
                    continue
                months = max(parsed_months)
                if 1 <= months <= 600:
                    score = 14 - lookahead
                    if "renewal term" in low_line:
                        score += 2
                    if "primary lease term" in low_line:
                        score += 3
                    candidates.append((score, idx, int(months)))
                    break
            year_match = re.search(r"(?i)\b(?:[a-z\-]+\s*)?\(?(\d{1,2})\)?\s*years?\b", probe)
            if year_match:
                years = _coerce_int_token(year_match.group(1), 0)
                if 1 <= years <= 50:
                    score = 12 - lookahead
                    if "renewal term" in low_line:
                        score += 2
                    candidates.append((score, idx, int(years * 12)))
                    break

    month_patterns = [
        (r"(?i)\b(?:primary\s+lease\s+term|lease\s+term|sublease\s+term|initial\s+term|term)\b[^.\n]{0,180}?\b\(?(\d{1,3})\)?\s*(?:calendar\s+)?months?\b", 9),
        (r"(?i)\b\(?(\d{1,3})\)?\s*(?:calendar\s+)?months?\s+(?:lease\s+|sublease\s+|initial\s+)?term\b", 11),
        (r"(?i)\bterm\s+and\s+free\s+rent\s*[:\-]\s*[^.\n]{0,160}?\((\d{1,3})\)\s*months?\b", 9),
        (r"(?i)\bterm\s+and\s+free\s+rent\s*[:\-]\s*[^.\n]{0,160}?\b(\d{1,3})\s*months?\b", 9),
        (r"(?i)\b(?:lease\s+term|sublease\s+term|initial\s+term|term\s*[:\-]|term\s+(?:shall\s+be|is|of))\b[^.\n]{0,160}?\((\d{1,3})\)\s*months?\b", 7),
        (r"(?i)\b(?:lease\s+term|sublease\s+term|initial\s+term|term\s*[:\-]|term\s+(?:shall\s+be|is|of))\b[^.\n]{0,160}?\b(\d{1,3})\s*months?\b", 6),
        (
            r"(?i)\b(?:landlord\s+proposes|proposes|proposal\s+is)\b[^\n]{0,120}?\b(?:[a-z\-]+\s*)?\(?(\d{1,3})\)?\s*"
            r"(?:calendar\s+)?months?\s+from\s+(?:the\s+)?lease\s+commencement(?:\s+date)?\b",
            13,
        ),
    ]
    for pattern, base_score in month_patterns:
        for match in re.finditer(pattern, text):
            if _is_disqualified(match):
                continue
            if _looks_like_non_term_month_context(_match_context(match)):
                continue
            months = _coerce_int_token(match.group(1), 0)
            if not (1 <= months <= 600):
                continue
            snippet = (match.group(0) or "").lower()
            score = base_score
            if "primary lease term" in snippet:
                score += 4
            elif re.search(r"(?i)\b(?:lease\s+term|sublease\s+term|initial\s+term)\b", snippet):
                score += 2
            if re.search(r"(?i)\bmonth\s+lease\s+term\b|\bmonths?\s+(?:lease\s+|sublease\s+|initial\s+)?term\b", snippet):
                score += 2
            if months >= 24:
                score += 1
            candidates.append((score, match.start(), int(months)))

    composite_patterns = [
        (r"(?i)\b(?:primary\s+lease\s+term|lease\s+term|sublease\s+term|initial\s+term|term)\b[^.\n]{0,180}?\b\(?(\d{1,2})\)?\s*years?\s*\+\s*\(?(\d{1,2})\)?\s*months?\b", 11),
        (r"(?i)\b\(?(\d{1,2})\)?\s*years?\s*\+\s*\(?(\d{1,2})\)?\s*months?\b[^.\n]{0,80}\b(?:lease\s+|sublease\s+|initial\s+)?term\b", 12),
    ]
    for pattern, base_score in composite_patterns:
        for match in re.finditer(pattern, text):
            if _is_disqualified(match):
                continue
            years = _coerce_int_token(match.group(1), 0)
            months = _coerce_int_token(match.group(2), 0)
            if not (1 <= years <= 50 and 0 <= months <= 11):
                continue
            total_months = years * 12 + months
            candidates.append((base_score, match.start(), int(total_months)))

    year_patterns = [
        (r"(?i)\b(?:primary\s+lease\s+term|lease\s+term|sublease\s+term|initial\s+term|term)\b[^.\n]{0,180}?\b\(?(\d{1,2})\)?\s*years?\b", 7),
        (r"(?i)\b\(?(\d{1,2})\)?\s*years?\s+(?:lease\s+|sublease\s+|initial\s+)?term\b", 9),
        (r"(?i)\b(?:lease\s+)?term\b[^.\n]{0,160}?\((\d{1,2})\)\s*years?\b", 6),
        (r"(?i)\b(?:lease\s+)?term\b[^.\n]{0,160}?\b(\d{1,2})\s*years?\b", 5),
    ]
    for pattern, base_score in year_patterns:
        for match in re.finditer(pattern, text):
            if _is_disqualified(match):
                continue
            years = _coerce_int_token(match.group(1), 0)
            if not (1 <= years <= 50):
                continue
            snippet = (match.group(0) or "").lower()
            score = base_score
            if "primary lease term" in snippet:
                score += 3
            candidates.append((score, match.start(), int(years * 12)))

    if candidates:
        candidates.sort(key=lambda row: (-row[0], -row[2], row[1]))
        return candidates[0][2]
    return None


def _convert_year_index_steps_to_month_steps(steps: list[dict[str, float | int]], one_based: bool) -> list[dict[str, float | int]]:
    offset = 1 if one_based else 0
    normalized_years: list[dict[str, float | int]] = []
    for step in sorted(steps, key=lambda s: (int(s["start"]), int(s["end"]))):
        year_start = max(0, int(step["start"]) - offset)
        year_end = max(year_start, int(step["end"]) - offset)
        normalized_years.append(
            {
                "start": year_start,
                "end": year_end,
                "rate_psf_yr": float(step["rate_psf_yr"]),
            }
        )
    if not normalized_years:
        return []
    # Some proposals use absolute lease-year labels (e.g. "Years 12-13").
    # Rebase to the first detected year so the extracted schedule always starts at month 0.
    min_year = min(int(step["start"]) for step in normalized_years)
    converted: list[dict[str, float | int]] = []
    for step in normalized_years:
        rebased_start = max(0, int(step["start"]) - min_year)
        rebased_end = max(rebased_start, int(step["end"]) - min_year)
        converted.append(
            {
                "start": rebased_start * 12,
                "end": ((rebased_end + 1) * 12) - 1,
                "rate_psf_yr": float(step["rate_psf_yr"]),
            }
        )
    return converted


def _looks_like_year_index_steps(steps: list[dict[str, float | int]], term_month_count: int) -> bool:
    """
    Detect common extraction mistakes where "Year 1, Year 2..." was interpreted
    as months 0,1,... instead of month ranges.
    """
    if len(steps) < 2:
        return False
    max_end = max(int(s["end"]) for s in steps)
    max_span = max((int(s["end"]) - int(s["start"]) + 1) for s in steps)
    # Strong signal: many tiny contiguous buckets like 0-0,1-1,2-2 (or 1-1,2-2...).
    if len(steps) >= 3 and max_end <= 12 and max_span <= 2:
        return True
    if term_month_count < 24:
        return False
    if max_end > 30 or max_span > 6:
        return False
    # If interpreted as months, schedule covers too little of the term.
    if (max_end + 1) >= max(18, term_month_count // 2):
        return False
    # If interpreted as years, coverage should be close to total term.
    years_coverage_months = (max_end + 1) * 12
    return years_coverage_months >= max(24, term_month_count - 24)


def _repair_and_extend_month_steps(
    steps: list[dict[str, float | int]], term_month_count: int
) -> list[dict[str, float | int]]:
    if not steps:
        return []
    ordered = sorted(steps, key=lambda s: (int(s["start"]), int(s["end"])))
    normalized: list[dict[str, float | int]] = []
    for raw in ordered:
        start = max(0, int(raw["start"]))
        end = max(start, int(raw["end"]))
        rate = max(0.0, float(raw["rate_psf_yr"]))
        if not normalized:
            if start != 0:
                start = 0
            normalized.append({"start": start, "end": end, "rate_psf_yr": rate})
            continue
        prev = normalized[-1]
        expected = int(prev["end"]) + 1
        if start > expected:
            prev["end"] = start - 1
        elif start < expected:
            start = expected
            if end < start:
                end = start
        normalized.append({"start": start, "end": end, "rate_psf_yr": rate})
    target_end = max(0, term_month_count - 1)
    if normalized[-1]["end"] < target_end:
        normalized[-1]["end"] = target_end
    return normalized


def _normalize_rent_steps(
    raw_steps: Any,
    term_month_count: int,
    prefill: dict | None = None,
) -> tuple[list[dict[str, float | int]], list[str]]:
    parsed: list[dict[str, float | int]] = []
    notes: list[str] = []
    fallback_rate = max(0.0, float(_coerce_float_token((prefill or {}).get("rate_psf_yr"), 0.0) or 0.0))
    for step in raw_steps or []:
        if not isinstance(step, dict):
            continue
        start = _coerce_int_token(step.get("start"), 0)
        end = _coerce_int_token(step.get("end"), start)
        rate = _coerce_float_token(step.get("rate_psf_yr"), fallback_rate)
        if start is None or end is None or rate is None:
            continue
        if end < start:
            end = start
        parsed.append({"start": start, "end": end, "rate_psf_yr": max(0.0, float(rate))})

    if not parsed:
        missing_rate_note = (
            "Rent steps missing; rate not found in document. Added placeholder step at $0.00 for review."
            if fallback_rate <= 0
            else "Rent steps missing; applied single step using extracted base rent."
        )
        return (
            [{"start": 0, "end": max(0, term_month_count - 1), "rate_psf_yr": fallback_rate}],
            [missing_rate_note],
        )

    basis = str((prefill or {}).get("_rent_steps_basis") or "").strip().lower()
    year_base_hint = _coerce_int_token((prefill or {}).get("_rent_steps_year_base"), None)
    looks_like_years = basis == "year_index" or _looks_like_year_index_steps(parsed, term_month_count)
    if looks_like_years:
        one_based = year_base_hint == 1
        if year_base_hint is None:
            min_year_index = min(min(int(s["start"]), int(s["end"])) for s in parsed)
            one_based = min_year_index >= 1
        parsed = _convert_year_index_steps_to_month_steps(parsed, one_based=one_based)
        notes.append(
            "Rent schedule interpreted as lease years and converted to month ranges "
            "(Year 1 = months 1-12, stored internally as 0-11)."
        )

    return _repair_and_extend_month_steps(parsed, term_month_count), notes


def _looks_like_fragmented_year_schedule(steps: list[dict[str, float | int]], term_month_count: int) -> bool:
    """
    Detect bad LLM outputs like:
      0-2, 3-3, 4-4, 5-5, 6-59
    which are typically lease-year labels misread as month buckets.
    """
    if len(steps) < 4:
        return False
    ordered = sorted(steps, key=lambda s: (int(s["start"]), int(s["end"])))
    if int(ordered[0]["start"]) != 0:
        return False
    if int(ordered[-1]["end"]) != max(0, term_month_count - 1):
        return False
    lead = ordered[:-1]
    if len(lead) < 3:
        return False
    max_lead_end = max(int(s["end"]) for s in lead)
    max_lead_span = max(int(s["end"]) - int(s["start"]) + 1 for s in lead)
    tail = ordered[-1]
    tail_span = int(tail["end"]) - int(tail["start"]) + 1
    if max_lead_end > 12 or max_lead_span > 3:
        return False
    if int(tail["start"]) > 12:
        return False
    if tail_span < max(18, term_month_count // 2):
        return False
    return True


def _extract_year_table_rent_steps(text: str) -> list[dict[str, float | int]]:
    """
    Parse common proposal tables like:
      Year 1: $45.00
      Years 2-3: $46.35
    and convert to month-index rent steps.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return []

    year_rows: list[tuple[int, int, float]] = []
    header_window = 0
    adjusted_window = 0
    scan_limit = min(len(lines), 1000)
    for idx, line in enumerate(lines[:scan_limit]):
        low_line = line.lower()
        if re.search(r"(?i)\badjusted\s+(?:rental|rent)\s+rate\b|\bgross\s+rent\b", line):
            adjusted_window = max(adjusted_window, 14)
        skip_adjusted_section = adjusted_window > 0 and bool(year_rows)
        if adjusted_window > 0:
            adjusted_window -= 1
        if skip_adjusted_section:
            continue
        for m in _RE_YEAR_RATE_INLINE_MONTHLY_PSF.finditer(line):
            y1 = _coerce_int_token(m.group(1), None)
            y2 = _coerce_int_token(m.group(2), y1)
            rate = _coerce_float_token(m.group(3), None)
            if y1 is None or y2 is None or rate is None:
                continue
            if not (2 <= rate <= 500):
                continue
            if y2 < y1:
                y1, y2 = y2, y1
            year_rows.append((y1, y2, float(rate)))
        if _RE_YEAR_TABLE_HEADER.search(line) or _RE_RENT_TABLE_HINT.search(line):
            header_window = max(header_window, 50)
        if header_window > 0:
            m = _RE_YEAR_TABLE_ROW.search(line)
            if m:
                y1 = _coerce_int_token(m.group(1), None)
                y2 = _coerce_int_token(m.group(2), y1)
                rate = _coerce_float_token(m.group(3), None)
                if y1 is not None and y2 is not None and rate is not None and 2 <= rate <= 500:
                    if y2 < y1:
                        y1, y2 = y2, y1
                    year_rows.append((y1, y2, float(rate)))
            else:
                label = _RE_YEAR_LABEL_ONLY.search(line)
                if label:
                    y1 = _coerce_int_token(label.group(1), None)
                    y2 = _coerce_int_token(label.group(2), y1)
                    if y1 is not None and y2 is not None:
                        if y2 < y1:
                            y1, y2 = y2, y1
                        for look_ahead in lines[idx + 1: idx + 5]:
                            rate_match = _RE_SIMPLE_RATE_LINE.search(look_ahead)
                            if not rate_match:
                                continue
                            rate = _coerce_float_token(rate_match.group(1), None)
                            if rate is None or not (2 <= rate <= 500):
                                continue
                            year_rows.append((y1, y2, float(rate)))
                            break
            header_window -= 1

        for m in _RE_YEAR_RATE_INLINE.finditer(line):
            if _RE_NON_RENT_YEAR_CONTEXT.search(low_line):
                continue
            y1 = _coerce_int_token(m.group(1), None)
            y2 = _coerce_int_token(m.group(2), y1)
            rate = _coerce_float_token(m.group(3), None)
            if y1 is None or y2 is None or rate is None:
                continue
            matched = line[m.start():m.end()]
            if "$" not in matched and not re.search(r"(?i)\b(?:rent|rate|psf|/sf|per\s+sf)\b", line):
                continue
            if not (2 <= rate <= 500):
                continue
            if y2 < y1:
                y1, y2 = y2, y1
            year_rows.append((y1, y2, float(rate)))

    if not year_rows:
        return []

    deduped: list[tuple[int, int, float]] = []
    seen: set[tuple[int, int, float]] = set()
    for y1, y2, rate in sorted(year_rows, key=lambda x: (x[0], x[1], x[2])):
        key = (int(y1), int(y2), round(float(rate), 4))
        if key in seen:
            continue
        seen.add(key)
        deduped.append((int(y1), int(y2), float(rate)))

    min_year = min(y1 for y1, _, _ in deduped)
    one_based = min_year >= 1
    raw_steps = [{"start": y1, "end": y2, "rate_psf_yr": rate} for y1, y2, rate in deduped]
    month_steps = _convert_year_index_steps_to_month_steps(raw_steps, one_based=one_based)
    if not month_steps:
        return []
    derived_term_months = max(int(step["end"]) for step in month_steps) + 1
    return _repair_and_extend_month_steps(month_steps, term_month_count=max(12, derived_term_months))


def _extract_month_phrase_rent_steps(text: str, term_month_count: int) -> list[dict[str, float | int]]:
    """
    Parse month-range proposal text like:
      "$38.00 per RSF ... for months 1-12 with 3% annual escalations starting in month 13"
    """
    if not text:
        return []
    base_match = _RE_MONTH_RANGE_RATE.search(text)
    if not base_match:
        return []
    rate = _coerce_float_token(base_match.group(1), None)
    start_month_1 = _coerce_int_token(base_match.group(2), None)
    end_month_1 = _coerce_int_token(base_match.group(3), None)
    if rate is None or start_month_1 is None or end_month_1 is None:
        return []
    if not (2 <= float(rate) <= 500):
        return []
    if start_month_1 <= 0 or end_month_1 < start_month_1:
        return []

    term = max(1, int(term_month_count))
    steps: list[dict[str, float | int]] = []
    first_start = max(1, start_month_1)
    first_end = min(term, max(first_start, end_month_1))
    steps.append(
        {
            "start": first_start - 1,
            "end": first_end - 1,
            "rate_psf_yr": round(float(rate), 2),
        }
    )

    esc_match = _RE_ANNUAL_ESCALATION.search(text)
    if not esc_match:
        if first_end < term:
            steps.append(
                {
                    "start": first_end,
                    "end": term - 1,
                    "rate_psf_yr": round(float(rate), 2),
                }
            )
        return _repair_and_extend_month_steps(steps, term_month_count=term)

    esc_pct = _coerce_float_token(esc_match.group(1), None)
    esc_start_1 = _coerce_int_token(esc_match.group(2), None)
    if esc_pct is None or esc_start_1 is None or esc_pct < 0 or esc_start_1 <= 0:
        return _repair_and_extend_month_steps(steps, term_month_count=term)
    esc_rate = float(esc_pct) / 100.0

    current_start_1 = max(esc_start_1, first_end + 1)
    current_rate = float(rate) * (1.0 + esc_rate)
    while current_start_1 <= term:
        current_end_1 = min(term, current_start_1 + 11)
        steps.append(
            {
                "start": current_start_1 - 1,
                "end": current_end_1 - 1,
                "rate_psf_yr": round(current_rate, 2),
            }
        )
        current_start_1 += 12
        current_rate *= 1.0 + esc_rate

    return _repair_and_extend_month_steps(steps, term_month_count=term)


def _infer_rate_psf_from_monthly_rent(text: str, rsf: float | None) -> float | None:
    if rsf is None or rsf <= 0:
        return None
    monthly_candidates: list[float] = []

    # Explicit monthly base rent amount.
    m = re.search(r"(?i)\bmonthly\s+(?:base\s+)?rent\b[^\n$]{0,40}\$?\s*([\d,]+(?:\.\d{1,2})?)", text)
    if m:
        v = _coerce_float_token(m.group(1), None)
        if v and 500 <= v <= 2_000_000:
            monthly_candidates.append(float(v))

    # "1 month of Prepaid Rent ($51,384.67)" pattern.
    m = re.search(r"(?i)\bprepaid\s+rent\b[^\n$]{0,80}\(\$?\s*([\d,]+(?:\.\d{1,2})?)\)", text)
    if m:
        v = _coerce_float_token(m.group(1), None)
        if v and 500 <= v <= 2_000_000:
            monthly_candidates.append(float(v))

    # "2 months' rent ($109,027.99)" => infer one month.
    m = re.search(
        r"(?i)\b(\d{1,2})\s+months?[’']?\s+rent\b[^\n$]{0,80}\(\$?\s*([\d,]+(?:\.\d{1,2})?)\)",
        text,
    )
    if m:
        months = _coerce_int_token(m.group(1), None)
        total = _coerce_float_token(m.group(2), None)
        if months and months > 0 and total and 500 <= total <= 4_000_000:
            monthly_candidates.append(float(total) / float(months))

    if not monthly_candidates:
        return None

    monthly = monthly_candidates[0]
    annual_psf = (monthly * 12.0) / float(rsf)
    if 2.0 <= annual_psf <= 500.0:
        return round(float(annual_psf), 4)
    return None


def _regex_prefill(text: str) -> dict:
    """Fast regex pass to prefill scenario fields. Returns dict of field -> value (or None)."""
    prefill = {}
    def _in_example_context(raw_text: str, idx: int) -> bool:
        prefix = raw_text[max(0, idx - 28):idx].lower()
        return ("e.g" in prefix) or ("example" in prefix)
    def _in_negotiation_timeline_context(raw_text: str, idx: int) -> bool:
        window = raw_text[max(0, idx - 90): min(len(raw_text), idx + 90)].lower()
        timeline_markers = (
            "ll proposal",
            "tenant response",
            "ll response",
            "landlord response",
            "counter proposal",
            "request for proposal",
            "lease proposal",
            "proposal prepared for",
            "delivered via",
            "via electronic mail",
        )
        if not any(marker in window for marker in timeline_markers):
            return False
        lease_markers = (
            "commencement",
            "expiration",
            "lease term",
            "premises",
            "base rent",
            "operating expenses",
        )
        return not any(marker in window for marker in lease_markers)

    # RSF: score all matches; avoid ratio/rate contexts (e.g., "per 1,000 RSF").
    rsf_best: tuple[int, int, float] | None = None  # (score, start, value)
    total_rsf_match = re.search(
        r"(?i)\btotal\s+premises\b\s*[:#-]?\s*(\d{1,3}(?:,\d{3})+|\d{3,7})\s*(?:rsf|sf|rentable\s+square\s+feet|square\s*feet)\b",
        text,
    )
    if total_rsf_match:
        total_value = _coerce_float_token(total_rsf_match.group(1), None)
        if total_value is not None and 100 <= total_value <= 2_000_000:
            rsf_best = (24, total_rsf_match.start(), float(total_value))
    for m in _RE_RSF.finditer(text):
        nums = _RE_NUM.findall(m.group(0))
        if not nums:
            continue
        try:
            value = float(nums[0].replace(",", ""))
        except ValueError:
            continue
        if not (100 <= value <= 2_000_000):
            continue
        start = max(0, m.start() - 140)
        end = min(len(text), m.end() + 140)
        snippet = text[start:end].replace("\n", " ").lower()
        score = 0
        if any(k in snippet for k in ("premises", "suite", "rentable", "square feet", "consisting of")):
            score += 4
        if any(k in snippet for k in ("spec suite", "lease premises", "premises:", "located in suite")):
            score += 4
        if 2_000 <= value <= 300_000:
            score += 1
        if "total premises" in snippet or "total rentable" in snippet:
            score += 8
        if any(k in snippet for k in ("project size", "building size", "between two buildings", "entire project")):
            score -= 7
        if re.search(r"(?i)\bper\s+1,?000\s+rsf\b|/1,?000\s*rsf|ratio", snippet):
            score -= 9
        if re.search(r"(?i)\(\s*\d[\d,]*\s*rsf\s*/\s*\d[\d,]*\s*rsf\s*\)", snippet):
            score -= 8
        if any(k in snippet for k in ("cam", "opex", "operating expenses", "taxes", "insurance", "parking")):
            score -= 4
        if any(k in snippet for k in ("shopping center contains", "total rsf", "entire center")):
            score -= 5
        if re.search(r"(?i)\$\s*\d[\d,]*(?:\.\d+)?\s*(?:/|per)\s*(?:sf|rsf)", snippet):
            score -= 4
        candidate = (score, m.start(), value)
        if rsf_best is None or candidate[0] > rsf_best[0] or (candidate[0] == rsf_best[0] and candidate[1] < rsf_best[1]):
            rsf_best = candidate
    if rsf_best is not None:
        prefill["rsf"] = rsf_best[2]
        prefill["_rsf_score"] = int(rsf_best[0])
    # Dates: prefer labeled commencement/expiration first, then generic first-two date fallback.
    comm = None
    exp = None
    m = _RE_COMM_LABEL.search(text)
    if m:
        comm = _parse_text_date_token(m.group(1))
    if not comm:
        m = re.search(r"(?i)\bcommenc(?:e|ing|ement)\b[^\n]{0,140}", text)
        if m and "e.g" not in (m.group(0) or "").lower():
            comm = _parse_text_date_token(m.group(0))
    m = _RE_EXP_LABEL.search(text)
    if m:
        exp = _parse_text_date_token(m.group(1))
    if not exp:
        m = re.search(r"(?i)\bending\s+on\b[^\n]{0,120}", text)
        if m and "e.g" not in (m.group(0) or "").lower():
            exp = _parse_text_date_token(m.group(0))
    if not exp:
        m = re.search(r"(?i)\b(?:sublease\s+term|term)\b[^\n]{0,200}\bthrough\b[^\n]{0,120}", text)
        if m and "e.g" not in (m.group(0) or "").lower():
            exp = _parse_text_date_token(m.group(0))
    if not comm or not exp:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for idx, ln in enumerate(lines):
            low_ln = ln.lower()
            is_comm_label = bool(
                "estimated commencement date" in low_ln
                or re.search(r"(?i)\bcommencement\s+date\s*[:\-]\s*$", low_ln)
            )
            is_exp_label = bool(
                "estimated termination date" in low_ln
                or "estimated expiration date" in low_ln
                or re.search(r"(?i)\b(?:expiration|termination)\s+date\s*[:\-]\s*$", low_ln)
            )
            if not comm and is_comm_label:
                for nxt in lines[idx + 1: idx + 8]:
                    if "e.g" in nxt.lower() or "example" in nxt.lower():
                        continue
                    parsed = _parse_text_date_token(nxt)
                    if parsed:
                        comm = parsed
                        break
            if not exp and is_exp_label:
                for nxt in lines[idx + 1: idx + 8]:
                    if "e.g" in nxt.lower() or "example" in nxt.lower():
                        continue
                    parsed = _parse_text_date_token(nxt)
                    if parsed:
                        exp = parsed
                        break
            if comm and exp:
                break
    if not comm or not exp:
        generic_dates: list[str] = []
        for hit in _RE_ISO_DATE.finditer(text):
            if _in_example_context(text, hit.start()):
                continue
            if _in_negotiation_timeline_context(text, hit.start()):
                continue
            iso = _parse_text_date_token(hit.group(0))
            if iso:
                generic_dates.append(iso)
        for hit in _RE_US_DATE.finditer(text):
            if _in_example_context(text, hit.start()):
                continue
            if _in_negotiation_timeline_context(text, hit.start()):
                continue
            iso = _parse_text_date_token(hit.group(0))
            if iso:
                generic_dates.append(iso)
        for hit in _RE_MONTH_DATE.finditer(text):
            if _in_example_context(text, hit.start()):
                continue
            if _in_negotiation_timeline_context(text, hit.start()):
                continue
            iso = _parse_text_date_token(hit.group(0))
            if iso:
                generic_dates.append(iso)
        # de-duplicate while preserving order
        seen = set()
        ordered = []
        for d in generic_dates:
            if d in seen:
                continue
            seen.add(d)
            ordered.append(d)
        if not comm and len(ordered) >= 1:
            comm = ordered[0]
        if not exp and ordered:
            if comm:
                comm_date = _safe_date(comm)
                for token in ordered:
                    parsed_token = _safe_date(token)
                    if parsed_token > comm_date:
                        exp = token
                        break
            elif len(ordered) >= 2:
                exp = ordered[1]

    def _derive_relative_commencement_iso(raw_text: str) -> str | None:
        clause_match = re.search(
            r"(?is)\b(?:lease\s+)?commencement(?:\s+date)?\s*[:\-]\s*([^\n]{0,280})",
            raw_text,
        )
        if not clause_match:
            return None
        clause = str(clause_match.group(1) or "").strip()
        low_clause = clause.lower()
        if not clause or ("following" not in low_clause and "after" not in low_clause):
            return None
        offset_match = re.search(r"(?i)\b(\d{1,3})\s+days?\s+(?:following|after)\b", clause)
        if not offset_match:
            return None
        day_offset = _coerce_int_token(offset_match.group(1), 0) or 0
        if day_offset <= 0 or day_offset > 365:
            return None
        date_token = (
            r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|"
            r"Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s*,?\s*\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4}"
            r"|\d{4}[-/]\d{1,2}[-/]\d{1,2}"
            r"|\d{1,2}[/-]\d{1,2}[/-]\d{4}"
            r"|\d{1,2}[/-]\d{1,2}[/-]\d{2}"
        )
        local_start = max(0, clause_match.start() - 120)
        local_end = min(len(raw_text), clause_match.end() + 900)
        local_scope = raw_text[local_start:local_end]
        anchor_patterns = [
            rf"(?is)\bdelivery\s+date(?:\s+of)?\b[^A-Za-z0-9]{{0,16}}({date_token})",
            rf"(?is)\bspec\s+suites?\b[^\n]{{0,220}}?\bdelivery\s+date\b[^\n]{{0,40}}({date_token})",
            rf"(?is)\bsubstantial\s+completion\b[^\n]{{0,200}}({date_token})",
        ]

        def _find_anchor(scope: str) -> date | None:
            for pat in anchor_patterns:
                m = re.search(pat, scope)
                if not m:
                    continue
                iso = _parse_text_date_token(m.group(1))
                if iso:
                    return _safe_date(iso)
            return None

        anchor = _find_anchor(local_scope) or _find_anchor(raw_text)
        if anchor is None:
            return None
        derived = anchor + timedelta(days=int(day_offset))
        if derived.day != 1:
            if derived.month == 12:
                derived = date(derived.year + 1, 1, 1)
            else:
                derived = date(derived.year, derived.month + 1, 1)
        return derived.isoformat()

    relative_comm = _derive_relative_commencement_iso(text)
    if relative_comm:
        comm = relative_comm
        if prefill.get("term_months") and not exp:
            exp = _expiration_from_term_months(_safe_date(relative_comm), int(prefill["term_months"])).isoformat()

    if comm:
        prefill["commencement"] = comm
    if exp:
        prefill["expiration"] = exp
    term_months = _extract_term_months_from_text(text)
    if term_months is not None:
        prefill["term_months"] = int(term_months)
        if comm and not exp:
            prefill["expiration"] = _expiration_from_term_months(_safe_date(comm), term_months).isoformat()
    elif comm and exp:
        prefill["term_months"] = _term_month_count(_safe_date(comm), _safe_date(exp))
    # TI allowance: capture both $/SF and total-dollar allowance language.
    ti_allowance_psf = None
    ti_allowance_total = None
    ti_psf_candidates: list[tuple[int, float, int]] = []
    ti_total_candidates: list[tuple[int, float, int]] = []
    psf_unit_pat = re.compile(
        rf"(?i)(?:/|per)?\s*(?:rentable\s+)?{_SF_UNIT_PATTERN}\b"
    )
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for idx, line in enumerate(lines):
        line_low = line.lower()
        line_has_ti_keyword = _has_ti_context_token(line)
        neighbor_window = " ".join(lines[max(0, idx - 1): idx + 2]).lower()
        line_is_allowance_with_ti_context = (
            "allowance" in line_low
            and (
                _has_ti_context_token(neighbor_window)
                or "tenant improvement" in neighbor_window
                or "tenant improvements" in neighbor_window
            )
        )
        if not line_has_ti_keyword and not line_is_allowance_with_ti_context:
            continue
        window = " ".join(lines[max(0, idx - 1): idx + 5])
        low_window = window.lower()
        if "e.g" in low_window or "example" in low_window:
            continue
        if "outside of the tenant improvement allowance" in line_low or "test-fit" in line_low or "test fit" in line_low:
            continue
        for amount_match in _RE_PSF_AMOUNT.finditer(window):
            local = window[max(0, amount_match.start() - 80): min(len(window), amount_match.end() + 36)]
            local_low = local.lower()
            if "$" not in (amount_match.group(0) or ""):
                continue
            if not re.search(r"(?i)\b(?:allowance|tenant improvements?|tia|ti)\b", local):
                continue
            if "test fit" in local_low or "test-fit" in local_low:
                continue
            if re.search(r"(?i)\b(?:base rent|rental rate)\b", local):
                continue
            try:
                value = float(amount_match.group(1).replace(",", ""))
            except ValueError:
                continue
            if not (0.5 <= value <= 500):
                if (
                    value > 500
                    and value <= 50_000
                    and psf_unit_pat.search(amount_match.group(0) or "")
                ):
                    scaled = value / 100.0
                    if 0.5 <= scaled <= 500:
                        value = scaled
                    else:
                        continue
                else:
                    continue
            score = 0
            if line_has_ti_keyword or line_is_allowance_with_ti_context:
                score += 6
            if any(
                tok in local_low
                for tok in (
                    "tenant improvement allowance",
                    "ti allowance",
                    "improvement allowance",
                    " tia ",
                )
            ):
                score += 8
            elif "allowance" in local_low:
                score += 4
            if re.search(r"(?i)\b(?:operating expenses?|opex|cam|base year)\b", local):
                score -= 10
            if "outside of the tenant improvement allowance" in local_low or "test-fit" in local_low or "test fit" in local_low:
                score -= 8
            if "not to exceed" in local_low:
                score += 8
            if re.search(r"(?i)\blandlord\s+shall\s+provide\b", local):
                score += 2
            ti_psf_candidates.append((score, value, idx))
        for total_match in re.finditer(r"\$\s*([\d,]+(?:\.\d+)?)", window):
            try:
                total_value = float(total_match.group(1).replace(",", ""))
            except ValueError:
                continue
            if total_value <= 500:
                continue
            local = window[max(0, total_match.start() - 80): min(len(window), total_match.end() + 90)]
            local_low = local.lower()
            if "outside of the tenant improvement allowance" in local_low:
                continue
            if psf_unit_pat.search(local):
                continue
            if not re.search(r"(?i)\b(?:allowance|tenant improvements?|tia|ti)\b", local):
                continue
            if "project management fee" in local_low:
                continue
            score = 0
            if any(tok in local_low for tok in ("tenant improvement allowance", "ti allowance", "improvement allowance")):
                score += 6
            if "allowance" in local_low:
                score += 3
            if any(tok in local_low for tok in ("budget", "cost", "fee")) and "allowance" not in local_low:
                score -= 6
            if total_value >= 1000:
                score += 1
            ti_total_candidates.append((score, float(total_value), idx))
    if ti_psf_candidates:
        ti_psf_candidates.sort(key=lambda row: (-row[0], -row[2], -row[1]))
        if ti_psf_candidates[0][0] > -2:
            ti_allowance_psf = ti_psf_candidates[0][1]
    if ti_allowance_psf is None:
        for m in _RE_TI.finditer(text):
            seg = (m.group(0) or "")
            seg_low = seg.lower()
            if "outside of the tenant improvement allowance" in seg_low:
                continue
            if "test fit" in seg_low or "test-fit" in seg_low:
                continue
            for amount_match in _RE_PSF_AMOUNT.finditer(seg):
                if "$" not in (amount_match.group(0) or ""):
                    continue
                fallback_val = _coerce_float_token(amount_match.group(1), 0.0) or 0.0
                if 0.5 <= fallback_val <= 500:
                    ti_allowance_psf = float(fallback_val)
                    break
            if ti_allowance_psf is not None:
                break
    if ti_total_candidates:
        ti_total_candidates.sort(key=lambda row: (-row[0], -row[2], -row[1]))
        if ti_total_candidates[0][0] > -2:
            ti_allowance_total = ti_total_candidates[0][1]
    if ti_allowance_total is None:
        total_candidates: list[tuple[int, float, int]] = []
        for m in _RE_TI_ALLOWANCE_TOTAL.finditer(text):
            value = _coerce_float_token(m.group(1), 0.0) or 0.0
            if value <= 500:
                continue
            seg = (m.group(0) or "").lower()
            context = text[max(0, m.start() - 80): min(len(text), m.end() + 120)]
            if psf_unit_pat.search(seg) or psf_unit_pat.search(context):
                continue
            score = 2
            if "tenant improvement allowance" in seg or "ti allowance" in seg:
                score += 4
            if "allowance" in seg:
                score += 2
            if any(tok in seg for tok in ("budget", "fee")) and "allowance" not in seg:
                score -= 5
            total_candidates.append((score, float(value), m.start()))
        if total_candidates:
            total_candidates.sort(key=lambda row: (-row[0], -row[2], -row[1]))
            if total_candidates[0][0] > -2:
                ti_allowance_total = total_candidates[0][1]
    if ti_allowance_total is not None:
        prefill["ti_allowance_total"] = float(round(ti_allowance_total, 2))
    if ti_allowance_psf is None and ti_allowance_total is not None:
        rsf_for_ti = _coerce_float_token(prefill.get("rsf"), None)
        if rsf_for_ti is not None and rsf_for_ti > 0:
            ti_allowance_psf = float(ti_allowance_total) / float(rsf_for_ti)
    if ti_allowance_psf is not None:
        prefill["ti_allowance_psf"] = float(round(ti_allowance_psf, 4))
    # Free rent months (line-aware; avoid renewal notice month counts).
    free_rent_candidates: list[tuple[int, int, int]] = []
    for line_idx, raw_line in enumerate([ln.strip() for ln in text.splitlines() if ln.strip()]):
        line = raw_line.lower()
        if not any(tok in line for tok in ("free rent", "abatement", "abated", "abate")):
            continue
        if any(tok in line for tok in ("renewal", "option", "notice", "no earlier than", "no later than", "prior to")):
            continue
        for m in re.finditer(
            r"(?i)\b(?:with\s+)?(?:the\s+first\s+)?(?:[a-z\-]+\s*\((\d{1,3})\)|\(?(\d{1,3})\)?)\s+months?\b",
            raw_line,
        ):
            local = raw_line[max(0, m.start() - 120): min(len(raw_line), m.end() + 20)].lower()
            if _looks_like_abatement_distribution_window(local):
                continue
            count = _coerce_int_token(m.group(1), 0) or _coerce_int_token(m.group(2), 0) or 0
            if count <= 0:
                continue
            score = 4
            if "term and free rent" in line:
                score += 4
            if "base free rent" in line:
                score += 2
            if count >= 60:
                score -= 6
            free_rent_candidates.append((score, line_idx, int(count)))
        m = _RE_FREE_RENT.search(raw_line)
        if m:
            local = raw_line[max(0, m.start() - 120): min(len(raw_line), m.end() + 20)].lower()
            if _looks_like_abatement_distribution_window(local):
                continue
            count = _coerce_int_token(m.group(1), 0) or 0
            if count > 0:
                score = 2
                if "term and free rent" in line:
                    score += 4
                free_rent_candidates.append((score, line_idx, int(count)))
    if free_rent_candidates:
        free_rent_candidates.sort(key=lambda row: (-row[0], row[1], row[2]))
        prefill["free_rent_months"] = int(free_rent_candidates[0][2])
    # Base opex $/sf
    m = _RE_BASE_OPEX.search(text)
    if m:
        seg = m.group(0)
        opex_candidates: list[float] = []
        for amount_match in _RE_PSF_AMOUNT.finditer(seg):
            try:
                value = float(amount_match.group(1).replace(",", ""))
            except ValueError:
                continue
            if 0 < value <= 150:
                opex_candidates.append(value)
        if not opex_candidates:
            amount_match = re.search(r"\$\s*([\d,]+(?:\.\d+)?)", seg)
            if amount_match:
                try:
                    value = float(amount_match.group(1).replace(",", ""))
                except ValueError:
                    value = 0.0
                if 0 < value <= 150:
                    opex_candidates.append(value)
        if opex_candidates:
            prefill["base_opex_psf_yr"] = float(sorted(opex_candidates)[0])
    # Base rent ($/sf/yr) — prefer explicit rent-labeled amounts and avoid OpEx/cam contamination.
    base_rate_candidates: list[tuple[int, float, int]] = []
    # Split-line/base-rent label fallback with strong priority.
    lines_for_rent = [ln.strip() for ln in text.splitlines() if ln.strip()]
    rent_label_pat = re.compile(
        r"(?i)\b(?:lease\s+rate|renewal\s+base\s+rent|initial\s+base\s+rent|base\s+rent|basic\s+rent|rental\s+rate|annual\s+rent)\b"
    )
    for idx, ln in enumerate(lines_for_rent[:1400]):
        if not rent_label_pat.search(ln):
            continue
        # Include a wider look-ahead window to capture split-line DOCX responses where
        # the rent value appears a few lines below the "Base Rent" label.
        window = " ".join(lines_for_rent[idx: idx + 8])
        for amount_match in re.finditer(
            r"(?i)\$\s*([\d,]+(?:\.\d{1,4})?)\s*(?:nnn|gross|full\s*service|modified\s*gross)?\b",
            window,
        ):
            value = _coerce_float_token(amount_match.group(1), None)
            if value is None or not (3 <= value <= 500):
                continue
            local = window[max(0, amount_match.start() - 80): min(len(window), amount_match.end() + 120)].lower()
            tight_local = window[max(0, amount_match.start() - 45): min(len(window), amount_match.end() + 45)].lower()
            if any(tok in tight_local for tok in ("operating expense", "opex", "cam", "base year", "per hour", "hourly", "hvac")):
                continue
            if any(tok in tight_local for tok in ("allowance", "amortiz", "interest", "tenant improvement", "ti allowance")):
                continue
            score = 12
            if "renewal base rent" in ln.lower():
                score += 4
            if "initial base rent" in ln.lower():
                score += 3
            if re.search(r"(?i)\b(?:nnn|gross|full\s*service|modified\s*gross)\b", local):
                score += 2
            if "annual increases" in local or "annual escalation" in local:
                score += 1
            if amount_match.start() > 220:
                score -= 6
            elif amount_match.start() > 140:
                score -= 3
            if value >= 15:
                score += 1
            elif value < 8:
                score -= 1
            # Inline month-table scaffolding often carries redline artifacts like "$4.50" vs "$47.50".
            if value < 10 and re.search(r"(?i)\bmonths?\s*1\s*(?:-|to|through|thru|–|—)\s*\d{1,3}\b", local):
                score -= 3
            base_rate_candidates.append((score, float(value), idx))

    for pat in (_RE_BASE_RENT_FLEX, _RE_BASE_RENT):
        for m in pat.finditer(text):
            try:
                v = float(m.group(1).replace(",", ""))
            except ValueError:
                continue
            if not (3 <= v <= 500):
                continue
            seg = (m.group(0) or "").lower()
            group_start = m.start(1) if m.lastindex and m.lastindex >= 1 else m.start()
            amount_local = text[max(0, group_start - 90): min(len(text), group_start + 90)].lower()
            score = 0
            if any(tok in seg for tok in ("/rsf", "/sf", "/psf", "per rsf", "per sf", "psf")):
                score += 5
            if "basic rent" in seg or "lease rate" in seg:
                score += 5
            if "initial base rent" in seg:
                score += 4
            if "base rent" in seg:
                score += 2
            if any(tok in seg for tok in ("annual increase", "annual escalation")):
                score += 1
            if any(tok in seg for tok in ("operating", "opex", "cam", "parking", "allowance", "ti allowance")):
                score -= 8
            if any(tok in amount_local for tok in ("operating expense", "opex", "cam", "base year")):
                score -= 10
            if "operating expenses" in seg and "renewal base rent" not in seg and "base rent:" not in seg and "lease rate" not in seg:
                continue
            if "in addition to base rent" in seg:
                continue
            if any(tok in seg for tok in ("abatement", "abated", "free rent")) and "basic rent" not in seg:
                score -= 6
            if "adjusted rental rate" in seg or "gross rent" in seg:
                score -= 6
            # Tie-break toward realistic base rent over low OpEx-like values when context score is similar.
            if v >= 15:
                score += 1
            elif v < 8:
                score -= 1
            base_rate_candidates.append((score, v, m.start()))
    if base_rate_candidates:
        base_rate_candidates.sort(key=lambda row: (-row[0], -row[1], row[2]))
        chosen = base_rate_candidates[0]
        best_score = int(chosen[0])
        low_candidates = [row for row in base_rate_candidates if row[1] < 10 and row[0] >= (best_score - 6)]
        high_candidates = [row for row in base_rate_candidates if row[1] >= 12 and row[0] >= (best_score - 6)]
        if low_candidates and high_candidates:
            best_high = sorted(high_candidates, key=lambda row: (-row[0], -row[1], row[2]))[0]
            if chosen[1] < 10 and best_high[1] >= max(12.0, chosen[1] * 2.0) and best_high[0] >= (chosen[0] - 4):
                chosen = best_high
            conflict_values = sorted(
                {
                    round(float(row[1]), 4)
                    for row in (low_candidates + high_candidates)
                }
            )
            if len(conflict_values) >= 2:
                prefill["_rate_psf_yr_conflict"] = "low_high_ambiguity"
                prefill["_rate_psf_yr_candidates"] = conflict_values[:8]
        prefill["rate_psf_yr"] = float(chosen[1])
    if "rate_psf_yr" not in prefill:
        fallback_candidates: list[tuple[int, float, int]] = []
        for idx, ln in enumerate(lines_for_rent[:1400]):
            if not rent_label_pat.search(ln):
                continue
            window = " ".join(lines_for_rent[idx: idx + 3])
            amount_match = re.search(
                r"(?i)\$\s*([\d,]+(?:\.\d{1,4})?)\s*(?:nnn|gross|full\s*service|modified\s*gross)?\b",
                window,
            )
            if not amount_match:
                continue
            low_window = window.lower()
            local = window[max(0, amount_match.start() - 60): min(len(window), amount_match.end() + 60)].lower()
            if any(tok in local for tok in ("operating expense", "opex", "cam")):
                continue
            value = _coerce_float_token(amount_match.group(1), None)
            if value is None or not (3 <= value <= 500):
                continue
            score = 7
            if "renewal base rent" in ln.lower():
                score += 3
            if "annual increases" in low_window or "annual escalation" in low_window:
                score += 2
            if re.search(r"(?i)\b(?:nnn|gross|full\s*service|modified\s*gross)\b", window):
                score += 2
            fallback_candidates.append((score, float(value), idx))
        if fallback_candidates:
            fallback_candidates.sort(key=lambda row: (-row[0], row[2], -row[1]))
            prefill["rate_psf_yr"] = float(fallback_candidates[0][1])
    if "rate_psf_yr" not in prefill:
        inferred_rate = _infer_rate_psf_from_monthly_rent(text, _coerce_float_token(prefill.get("rsf"), None))
        if inferred_rate is not None:
            prefill["rate_psf_yr"] = inferred_rate
    month_phrase_steps: list[dict[str, float | int]] = []
    if term_months:
        month_phrase_steps = _extract_month_phrase_rent_steps(text, term_month_count=int(term_months))
    if month_phrase_steps:
        prefill["rent_steps"] = month_phrase_steps
        prefill["_rent_steps_basis"] = "month_index"
        prefill["_rent_steps_source"] = "month_phrase_regex"
        prefill["rate_psf_yr"] = float(month_phrase_steps[0].get("rate_psf_yr", prefill.get("rate_psf_yr", 0.0)))
    year_table_steps = _extract_year_table_rent_steps(text)
    if year_table_steps and "rent_steps" not in prefill:
        inferred_term_from_dates = None
        if comm and exp:
            inferred_term_from_dates = _term_month_count(_safe_date(comm), _safe_date(exp))
        year_table_term = max(int(step["end"]) for step in year_table_steps) + 1
        # Guardrail: in amendment/sublease files, master-lease year tables can appear
        # and be wildly out of range for the current document's term.
        if not (
            inferred_term_from_dates
            and inferred_term_from_dates <= 36
            and year_table_term > inferred_term_from_dates * 4
        ):
            prefill["rent_steps"] = year_table_steps
            prefill["_rent_steps_basis"] = "month_index"
            prefill["_rent_steps_source"] = "year_table_regex"
            prefill["rate_psf_yr"] = float(year_table_steps[0].get("rate_psf_yr", prefill.get("rate_psf_yr", 0.0)))

    # If only base rate + annual escalation language exists, synthesize month-index steps.
    # Respect explicit escalation start months (e.g., "escalations beginning month 19").
    if "rate_psf_yr" in prefill:
        existing_steps = prefill.get("rent_steps") if isinstance(prefill.get("rent_steps"), list) else []
        existing_source = str(prefill.get("_rent_steps_source") or "").strip().lower()
        esc_rate = _extract_annual_rent_escalation_pct(text)
        esc_start_month_1 = _extract_annual_rent_escalation_start_month(text)
        inferred_term = _coerce_int_token(prefill.get("term_months"), None)
        if inferred_term is None and comm and exp:
            inferred_term = _term_month_count(_safe_date(comm), _safe_date(exp))
        existing_term = (
            max((int(step.get("end", 0)) for step in existing_steps if isinstance(step, dict)), default=-1) + 1
            if existing_steps
            else 0
        )
        existing_rate_count = len(
            {
                round(_coerce_float_token(step.get("rate_psf_yr"), 0.0) or 0.0, 4)
                for step in existing_steps
                if isinstance(step, dict)
            }
        )
        weak_flat_year_table_steps = (
            bool(existing_steps)
            and existing_source == "year_table_regex"
            and existing_rate_count <= 1
            and (inferred_term is None or existing_term < int(inferred_term))
        )
        should_synthesize = (not existing_steps) or weak_flat_year_table_steps
        if (
            should_synthesize
            and
            esc_rate is not None
            and inferred_term is not None
            and inferred_term > 1
            and 0 < float(esc_rate) <= 25.0
        ):
            annual = float(esc_rate) / 100.0
            base_rate = float(prefill["rate_psf_yr"])
            term = int(inferred_term)
            synth_steps: list[dict[str, float | int]] = []
            start_month_1 = max(1, int(esc_start_month_1 or 13))
            if start_month_1 > term:
                synth_steps.append(
                    {
                        "start": 0,
                        "end": term - 1,
                        "rate_psf_yr": round(base_rate, 2),
                    }
                )
            else:
                # Pre-escalation base period.
                if start_month_1 > 1:
                    synth_steps.append(
                        {
                            "start": 0,
                            "end": max(0, start_month_1 - 2),
                            "rate_psf_yr": round(base_rate, 2),
                        }
                    )
                # Escalated periods roll every 12 months from the explicit start month.
                current_start_1 = start_month_1
                escalation_idx = 1
                while current_start_1 <= term:
                    current_end_1 = min(term, current_start_1 + 11)
                    current_rate = base_rate * ((1.0 + annual) ** escalation_idx)
                    synth_steps.append(
                        {
                            "start": current_start_1 - 1,
                            "end": current_end_1 - 1,
                            "rate_psf_yr": round(current_rate, 2),
                        }
                    )
                    current_start_1 += 12
                    escalation_idx += 1
            if synth_steps:
                prefill["rent_steps"] = synth_steps
                prefill["_rent_steps_basis"] = "month_index"
                prefill["_rent_steps_source"] = "base_rate_plus_escalation_regex"
    # Opex mode
    detected_opex_mode = _detect_opex_mode_from_text(text)
    if detected_opex_mode:
        prefill["opex_mode"] = detected_opex_mode
    # Scenario name from building/suite if possible
    building = None
    suite = None
    building_row_match = re.search(r"(?im)^\s*building\s*:\s*\|\s*([^|\n]+)", text)
    if not building_row_match:
        building_row_match = re.search(r"(?im)^\s*building\s*:\s*([^\n|]+)", text)
    if building_row_match:
        building = " ".join((building_row_match.group(1) or "").split()).strip(" ,.-")
        building = re.sub(
            r"(?i)\s+[–-]\s+\d{1,6}\s+[A-Za-z0-9 .,'-]{2,120}\b(?:street|st\.?|avenue|ave\.?|boulevard|blvd\.?|road|rd\.?|drive|dr\.?|lane|ln\.?|way|plaza|parkway|pkwy\.?|expressway|expy\.?|highway|hwy\.?)\b.*$",
            "",
            building,
        ).strip(" ,.-")
    premises_line_match = re.search(r"(?im)^\s*premises\s*:\s*([^\n]+)$", text)
    if premises_line_match:
        premises_line = premises_line_match.group(1)
        m = re.search(r"(?i)\b(?:suite|ste\.?|unit)\s*([A-Za-z0-9][A-Za-z0-9\-]{0,20})\b", premises_line)
        if m:
            suite = m.group(1).strip(" ,.-")
    m = _RE_BUILDING.search(text)
    if m and not building:
        building = m.group(1).strip(" ,.-")
    if not building:
        m = re.search(
            r"(?i)\blease\s+(?:space|premises)\s+at\s+([A-Za-z0-9&' .-]{3,80}?)(?:\s*[–-]\s*building\s*([A-Za-z0-9]+))?\b",
            text,
        )
        if m:
            bname = m.group(1).strip(" ,.-")
            bnum = (m.group(2) or "").strip()
            building = f"{bname} - Building {bnum}".strip(" -") if bnum else bname
    if suite is None:
        suite_match = _RE_SUITE.search(text)
        if suite_match:
            suite = suite_match.group(1).strip(" ,.-")
    if building and suite:
        prefill["name"] = f"{building} Suite {suite}"
    elif building:
        prefill["name"] = building
    elif suite:
        prefill["name"] = f"Suite {suite}"

    # Parking economics used by deterministic fallback when AI is unavailable.
    parking_ratio_evidence = False
    parking_count_evidence = False
    parking_ratio_patterns = [
        r"(?i)\b(\d+(?:\.\d+)?)\s*(?:(?:reserved|unreserved|covered|surface|garage)\s+){0,2}(?:parking\s+)?(?:spaces?|stalls?|passes?)\s*(?:per|\/)\s*(?:every\s*)?1,?000\s*(?:(?:rentable\s+)?square\s*feet|lease\s+space|rsf|sf|spaces?)\b",
        r"(?i)\b(\d+(?:\.\d+)?)\s*(?:/|per)\s*(?:every\s*)?1,?000\s*(?:(?:rentable\s+)?square\s*feet|lease\s+space|rsf|sf|spaces?)\b",
        r"(?i)\bparking\s+ratio\b[^.\n]{0,100}\b(\d+(?:\.\d+)?)\s*(?:/|per|:)\s*(?:every\s*)?1,?000\b",
        r"(?i)\bparking\s+ratio\s*(?:/|per)?\s*(?:every\s*)?1,?000\s*(?:rsf|sf|(?:rentable\s+)?square\s*feet)?\s*[:\-]?\s*(\d+(?:\.\d+)?)\b",
    ]
    parking_ratio_candidates: list[tuple[int, int, float]] = []
    for pat in parking_ratio_patterns:
        for m in re.finditer(pat, text):
            ratio = _coerce_float_token(m.group(1), 0.0) or 0.0
            if not (0.1 <= ratio <= 30):
                continue
            local = text[max(0, m.start() - 90): min(len(text), m.end() + 90)].lower()
            score = 1
            if any(k in local for k in ("parking", "ratio", "spaces", "garage", "permits")):
                score += 4
            if any(k in local for k in ("density", "desk", "workstation", "occupancy")):
                score -= 6
            parking_ratio_candidates.append((score, m.start(), float(ratio)))
    if parking_ratio_candidates:
        parking_ratio_candidates.sort(key=lambda row: (-row[0], -row[1], row[2]))
        if parking_ratio_candidates[0][0] > 0:
            prefill["parking_ratio_per_1000_rsf"] = float(parking_ratio_candidates[0][2])
            parking_ratio_evidence = True

    parking_count_patterns = [
        r"(?i)\bparking\s+spaces?\s*[:\-]?\s*(\d{1,4})(?!\.\d)\b",
        r"(?i)(?<![\d.])(\d{1,4})(?!\.\d)\s+(?:reserved|unreserved|covered|surface|garage)?\s*parking\s+spaces?\b",
        r"(?i)\b(?:up to\s+)?[a-z][a-z\-]*\s*\((\d{1,4})(?!\.\d)\)\s+(?:reserved|unreserved|covered|surface|garage)?\s*parking\s+spaces?\b",
    ]
    parking_count_candidates: list[tuple[int, int, int]] = []
    for pat in parking_count_patterns:
        for m in re.finditer(pat, text):
            count = _coerce_int_token(m.group(1), 0)
            if not (count and 1 <= count <= 10000):
                continue
            # Ignore decimal tails and ratio fragments such as ".32 parking spaces per every 1,000 ...".
            if m.start() > 0 and text[m.start() - 1] == ".":
                continue
            trailing = text[m.end(): min(len(text), m.end() + 64)].lower()
            if re.search(r"(?i)\b(?:per|\/)\s*(?:every\s*)?1,?000\b", trailing):
                continue
            local = text[max(0, m.start() - 90): min(len(text), m.end() + 110)].lower()
            score = 1
            if any(tok in local for tok in ("entitled", "allotted", "total # paid spaces", "total paid spaces", "# reserved", "# unreserved")):
                score += 4
            if any(tok in local for tok in ("reserved", "unreserved", "garage", "surface", "covered")):
                score += 2
            parking_count_candidates.append((score, m.start(), int(count)))
    if parking_count_candidates:
        parking_count_candidates.sort(key=lambda row: (-row[0], row[1], row[2]))
        if parking_count_candidates[0][0] > 0:
            prefill["parking_spaces"] = int(parking_count_candidates[0][2])
            parking_count_evidence = True
    reserved_count_match = re.search(r"(?i)\b#?\s*reserved\s+(?:paid\s+)?spaces?\b\s*[:\-]?\s*(\d{1,4}(?:\.\d+)?)", text)
    unreserved_count_match = re.search(r"(?i)\b#?\s*unreserved\s+(?:paid\s+)?spaces?\b\s*[:\-]?\s*(\d{1,4}(?:\.\d+)?)", text)
    total_paid_spaces_match = re.search(r"(?i)\btotal\s*#?\s*paid\s+spaces?\b\s*[:\-]?\s*(\d{1,4}(?:\.\d+)?)", text)
    reserved_count = _coerce_int_token(reserved_count_match.group(1), 0) if reserved_count_match else 0
    unreserved_count = _coerce_int_token(unreserved_count_match.group(1), 0) if unreserved_count_match else 0
    if reserved_count or unreserved_count:
        prefill["parking_spaces"] = max(0, int(reserved_count or 0)) + max(0, int(unreserved_count or 0))
        parking_count_evidence = True
    elif total_paid_spaces_match:
        total_paid_spaces = _coerce_int_token(total_paid_spaces_match.group(1), 0) or 0
        if total_paid_spaces > 0:
            prefill["parking_spaces"] = int(total_paid_spaces)
            parking_count_evidence = True

    rsf_score_for_parking = int(_coerce_int_token(prefill.get("_rsf_score"), -999) or -999)
    rsf_for_parking_alignment = prefill.get("rsf") if rsf_score_for_parking >= -8 else None
    aligned_ratio, aligned_count = _reconcile_parking_ratio_and_count(
        parking_ratio=prefill.get("parking_ratio_per_1000_rsf"),
        parking_count=prefill.get("parking_spaces"),
        rsf=rsf_for_parking_alignment,
        prefer="count" if parking_count_evidence and not parking_ratio_evidence else "ratio",
    )
    if aligned_ratio > 0:
        prefill["parking_ratio_per_1000_rsf"] = float(aligned_ratio)
    if aligned_count > 0:
        prefill["parking_spaces"] = int(aligned_count)

    parking_rate_patterns = [
        r"(?i)\$\s*([\d,]+(?:\.\d{1,2})?)\s*(?:per|\/)\s*(?:space|stall)\s*(?:per|\/)\s*month\b",
        r"(?i)\bparking\b[^.\n]{0,120}\$\s*([\d,]+(?:\.\d{1,2})?)\s*(?:per|\/)\s*(?:month|mo\.?)\b",
    ]
    for pat in parking_rate_patterns:
        m = re.search(pat, text)
        if not m:
            continue
        rate = _coerce_float_token(m.group(1), 0.0)
        if rate and 1 <= rate <= 10000:
            prefill["parking_cost_monthly_per_space"] = float(rate)
            break
    return prefill


def extract_prefill_hints(text: str) -> dict:
    """Public wrapper used by main normalizer to reuse deterministic regex prefill extraction."""
    return _regex_prefill(text or "")


def _safe_date(s: str | None) -> date:
    if not s:
        return date(2026, 1, 1)
    try:
        return date.fromisoformat(s.strip()[:10])
    except (ValueError, TypeError):
        return date(2026, 1, 1)


def _llm_extract_scenario(text: str, prefill: dict) -> dict:
    """Call LLM to fill gaps and validate; prefill contains regex-extracted values."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key or not str(api_key).strip():
        raise ValueError("OPENAI_API_KEY not configured")
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai required: pip install openai")
    timeout_sec = 60.0
    try:
        timeout_env = float((os.environ.get("OPENAI_EXTRACT_TIMEOUT_SEC") or "60").strip())
        if timeout_env > 0:
            timeout_sec = timeout_env
    except (TypeError, ValueError, AttributeError):
        pass

    prefill_for_model = {k: v for k, v in (prefill or {}).items() if not str(k).startswith("_")}
    prefill_hint = ""
    if prefill_for_model:
        prefill_hint = f"\nPre-extracted values (use these when confident, else null + warning): {json.dumps(prefill_for_model)}"

    prompt = f"""Extract lease terms as JSON only. No markdown, no prose.
- Output a single JSON object: {{ "scenario": {{ ... }}, "confidence": {{ "rsf": 0.9, ... }}, "warnings": [] }}
- Scenario fields to fill when present: name, rsf, commencement, expiration, rent_steps, free_rent_months, ti_allowance_psf, opex_mode, base_opex_psf_yr, base_year_opex_psf_yr, opex_growth.
- Treat lease year tables correctly: Year 1 = months 1-12 (stored internally as 0-11), Year 2 = months 13-24, etc.
- rent_steps must be month-index ranges only: [{{"start": 0, "end": 59, "rate_psf_yr": number}}, ...]
- rent_steps must be contiguous, no overlaps/gaps, sorted ascending.
- Prefer the subject Premises/Suite/Space for RSF and dates; ignore ROFR/ROFO/expansion/option spaces unless they are explicitly the leased premises.
- Do not invent values. Use null and add a warning when unknown.
- If only monthly rent and RSF are provided, infer annual PSF rent.
{prefill_hint}

Text:
---
{text}
---

JSON:"""

    configured = (os.environ.get("OPENAI_LEASE_MODEL") or "").strip()
    models = (
        [m.strip() for m in configured.split(",") if m.strip()]
        if configured
        else ["gpt-4o-mini", "gpt-4.1-mini", "gpt-4.1"]
    )
    last_error: Exception | None = None
    # Allow multiple model attempts before falling back to deterministic extraction.
    deadline = time.monotonic() + max(20.0, timeout_sec * max(1, min(len(models), 2)))
    for model in models:
        remaining = deadline - time.monotonic()
        if remaining <= 1.0:
            break
        per_call_timeout = max(10.0, min(timeout_sec, remaining))
        try:
            client = OpenAI(api_key=api_key, timeout=per_call_timeout)
            t0 = time.perf_counter()
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            elapsed = time.perf_counter() - t0
            logger.info("[extract] LLM call duration=%.2fs model=%s", elapsed, model)

            raw = response.choices[0].message.content.strip()
            if raw.startswith("```"):
                raw = re.sub(r"^```\w*\n?", "", raw)
                raw = re.sub(r"\n?```\s*$", "", raw)
            out = json.loads(raw)
            if not isinstance(out, dict):
                raise ValueError("LLM output was not a JSON object.")
            # Merge prefill into scenario so LLM doesn't have to repeat
            scenario = out.get("scenario") if isinstance(out.get("scenario"), dict) else {}
            for k, v in prefill_for_model.items():
                if v is not None and (scenario.get(k) is None or scenario.get(k) == ""):
                    scenario[k] = v
            out["scenario"] = scenario
            return out
        except Exception as e:
            last_error = e
            logger.warning("[extract] model failed model=%s error=%s", model, e)
            low = str(e).lower()
            # Fail fast only for hard auth/config errors.
            if any(x in low for x in ("invalid api key", "incorrect api key", "authentication", "unauthorized", "api key")):
                break
            continue
    if last_error is not None:
        raise last_error
    raise RuntimeError("No OpenAI model candidates configured")


def _infer_scenario_name(text: str) -> str:
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    labeled_patterns = [
        re.compile(r"(?i)\b(?:building|property)\s*(?:name)?\s*[:#-]\s*([^\n,;]{3,80})"),
        re.compile(r"(?i)\bpremises\s*[:#-]\s*(?:suite|ste\.?|unit)?\s*([A-Za-z0-9][A-Za-z0-9\- ]{1,40})"),
    ]
    for ln in lines[:80]:
        for pat in labeled_patterns:
            m = pat.search(ln)
            if not m:
                continue
            candidate = " ".join(m.group(1).split()).strip(" ,.;:-")
            if candidate and 2 <= len(candidate) <= 80:
                return candidate

    banned_fragments = (
        "hereby leases",
        "landlord",
        "tenant",
        "this lease",
        "agreement made",
        "commencement",
        "expiration",
        "term",
        "rent",
        "operating expenses",
        "table of contents",
    )
    for ln in lines[:40]:
        if len(ln) > 90:
            continue
        # Skip headings that are usually generic
        low = ln.lower()
        if any(k in low for k in ("lease", "agreement", "exhibit", "table of contents", "page ")):
            continue
        if any(k in low for k in banned_fragments):
            continue
        if re.fullmatch(r"[\d\W_]+", ln):
            continue
        # Avoid picking long legal clauses as scenario names.
        if "," in ln and len(ln.split()) > 9:
            continue
        if re.search(r"[A-Za-z]", ln):
            return ln
    return "Extracted lease"


def _heuristic_extract_scenario(text: str, prefill: dict, llm_error: Exception | None = None) -> tuple[dict, dict[str, float], list[str]]:
    """
    Deterministic fallback when LLM extraction is unavailable.
    Returns scenario dict compatible with Scenario model + confidence + warnings.
    """
    opex_mode = _detect_opex_mode_from_text(text) or "nnn"
    rsf = float(prefill.get("rsf", 10000.0) or 10000.0)
    rate = float(prefill.get("rate_psf_yr", 0.0) or 0.0)
    commencement = str(prefill.get("commencement") or "2026-01-01")
    expiration = str(prefill.get("expiration") or "2031-01-31")
    term_month_hint = _coerce_int_token(prefill.get("term_months"), None)
    term_month_count = int(term_month_hint) if term_month_hint and term_month_hint > 0 else _term_month_count(_safe_date(commencement), _safe_date(expiration))
    prefill_steps = prefill.get("rent_steps") if isinstance(prefill.get("rent_steps"), list) else None
    if prefill_steps:
        rent_steps, rent_step_notes = _normalize_rent_steps(prefill_steps, term_month_count=term_month_count, prefill=prefill)
    else:
        rent_steps, rent_step_notes = _normalize_rent_steps(
            [{"start": 0, "end": max(0, term_month_count - 1), "rate_psf_yr": rate}],
            term_month_count=term_month_count,
            prefill=prefill,
        )

    scenario = {
        "name": str(prefill.get("name") or _infer_scenario_name(text)),
        "rsf": rsf,
        "commencement": commencement,
        "expiration": expiration,
        "rent_steps": rent_steps,
        "free_rent_months": int(prefill.get("free_rent_months", 0) or 0),
        "ti_allowance_psf": float(prefill.get("ti_allowance_psf", 0.0) or 0.0),
        "opex_mode": opex_mode,
        "base_opex_psf_yr": float(prefill.get("base_opex_psf_yr", 10.0) or 10.0),
        "base_year_opex_psf_yr": float(prefill.get("base_opex_psf_yr", 10.0) or 10.0),
        "opex_growth": 0.03,
        "discount_rate_annual": 0.08,
        "parking_spaces": int(prefill.get("parking_spaces", 0) or 0),
        "parking_cost_monthly_per_space": float(prefill.get("parking_cost_monthly_per_space", 0.0) or 0.0),
        "parking_sales_tax_rate": 0.0825,
    }

    confidence = {
        "rsf": 0.7 if "rsf" in prefill else 0.0,
        "commencement": 0.6 if "commencement" in prefill else 0.0,
        "expiration": 0.6 if "expiration" in prefill else 0.0,
        "rent_steps": 0.7 if "rent_steps" in prefill else (0.5 if "rate_psf_yr" in prefill else 0.2),
        "ti_allowance_psf": 0.6 if "ti_allowance_psf" in prefill else 0.0,
        "free_rent_months": 0.6 if "free_rent_months" in prefill else 0.0,
    }
    warnings: list[str] = []
    warnings.extend(rent_step_notes)
    if llm_error:
        warnings.append("Automatic extraction fallback was used for this upload.")
        msg = str(llm_error).lower()
        if "openai_api_key" in msg:
            warnings.append("OPENAI_API_KEY is not configured on backend.")
        elif "invalid api key" in msg or "incorrect api key" in msg or "unauthorized" in msg or "authentication" in msg:
            warnings.append("OPENAI_API_KEY is invalid for this backend service.")
        elif "model" in msg and ("not found" in msg or "does not exist" in msg or "not have access" in msg):
            warnings.append("Configured OpenAI model is unavailable for this API key.")
        elif "rate" in msg or "quota" in msg or "429" in msg:
            warnings.append("AI provider rate limit/quota reached.")
        elif "timeout" in msg:
            warnings.append("AI extraction timed out.")
        elif "connection" in msg or "network" in msg or "api connection" in msg:
            warnings.append("Backend could not connect to OpenAI API.")
    return scenario, confidence, warnings


def _apply_safe_defaults(raw: dict | Any, prefill: dict | None = None) -> tuple[dict, dict[str, float], list[str]]:
    """Apply safe defaults to raw LLM output; return (scenario_dict, confidence, warnings)."""
    if not isinstance(raw, dict):
        raw = {"scenario": {}, "confidence": {}, "warnings": ["AI response shape was invalid; used fallback defaults."]}

    scenario = raw.get("scenario") if isinstance(raw.get("scenario"), dict) else {}
    confidence_raw = raw.get("confidence")
    confidence = dict(confidence_raw) if isinstance(confidence_raw, dict) else {}
    warnings_raw = raw.get("warnings")
    warnings = list(warnings_raw) if isinstance(warnings_raw, list) else []

    name = scenario.get("name")
    if not name or not str(name).strip():
        name = "Extracted lease"
        warnings.append("Scenario name missing; default applied.")
    scenario["name"] = str(name).strip()

    rsf = scenario.get("rsf")
    if rsf is None or (isinstance(rsf, (int, float)) and (rsf <= 0 or rsf != rsf)):
        scenario["rsf"] = 10000.0
        if "rsf" not in confidence:
            confidence["rsf"] = 0.0
        warnings.append("RSF missing or invalid; default 10,000 applied.")
    else:
        scenario["rsf"] = float(rsf)

    term_months_hint = _coerce_int_token(
        scenario.get("term_months"),
        _coerce_int_token((prefill or {}).get("term_months"), None),
    )
    commencement = scenario.get("commencement")
    commencement_date = commencement if isinstance(commencement, date) else _safe_date(commencement if isinstance(commencement, str) else None)
    scenario["commencement"] = commencement_date
    if not commencement:
        warnings.append("Commencement date missing; default 2026-01-01 applied.")
        confidence.setdefault("commencement", 0.0)

    expiration = scenario.get("expiration")
    expiration_date = expiration if isinstance(expiration, date) else _safe_date(expiration if isinstance(expiration, str) else None)
    expiration_inferred_from_term = False
    if term_months_hint and term_months_hint > 1:
        is_missing_exp = not expiration
        is_degenerate_exp = expiration_date <= commencement_date
        if is_missing_exp or is_degenerate_exp:
            expiration_date = _expiration_from_term_months(commencement_date, int(term_months_hint))
            expiration_inferred_from_term = True
            if is_missing_exp:
                warnings.append(f"Expiration date inferred from commencement + {int(term_months_hint)} month term.")
            else:
                warnings.append(f"Expiration date corrected from {int(term_months_hint)} month lease term.")
    scenario["expiration"] = expiration_date
    if not expiration and not expiration_inferred_from_term:
        warnings.append("Expiration date missing; default 2031-01-31 applied.")
        confidence.setdefault("expiration", 0.0)

    term_month_count = int(term_months_hint) if term_months_hint and term_months_hint > 0 else _term_month_count(commencement_date, expiration_date)
    rent_steps_input = scenario.get("rent_steps")
    had_explicit_rent_steps = isinstance(rent_steps_input, list) and len(rent_steps_input) > 0
    prefill_steps = (prefill or {}).get("rent_steps") if isinstance((prefill or {}).get("rent_steps"), list) else []
    if not isinstance(rent_steps_input, list):
        rent_steps_input = []
    if len(rent_steps_input) == 0 and prefill_steps:
        rent_steps_input = prefill_steps
        warnings.append("Rent steps inferred from rent schedule table in the uploaded file.")
    elif len(rent_steps_input) > 0 and prefill_steps:
        parsed_explicit: list[dict[str, float | int]] = []
        for step in rent_steps_input:
            if not isinstance(step, dict):
                continue
            start = _coerce_int_token(step.get("start"), None)
            end = _coerce_int_token(step.get("end"), start)
            rate = _coerce_float_token(step.get("rate_psf_yr"), None)
            if start is None or end is None or rate is None:
                continue
            parsed_explicit.append({"start": int(start), "end": int(max(start, end)), "rate_psf_yr": float(rate)})
        if parsed_explicit and _looks_like_fragmented_year_schedule(parsed_explicit, term_month_count):
            rent_steps_input = prefill_steps
            warnings.append(
                "Rent steps from AI output looked fragmented; replaced with rent schedule table extraction "
                    "(Year 1 = months 1-12, stored internally as 0-11)."
            )

    normalized_steps, rent_step_notes = _normalize_rent_steps(
        rent_steps_input,
        term_month_count=term_month_count,
        prefill=prefill,
    )
    scenario["rent_steps"] = normalized_steps
    warnings.extend(rent_step_notes)
    if not had_explicit_rent_steps:
        confidence.setdefault("rent_steps", 0.0)

    scenario.setdefault("free_rent_months", 0)
    scenario.setdefault("ti_allowance_psf", 0.0)
    opex = scenario.get("opex_mode")
    if opex not in ("nnn", "base_year", "full_service"):
        inferred_opex = str((prefill or {}).get("opex_mode") or "").strip().lower()
        if inferred_opex in {"nnn", "base_year", "full_service"}:
            scenario["opex_mode"] = inferred_opex
            warnings.append("Opex mode inferred from lease structure language in the document.")
        else:
            scenario["opex_mode"] = "nnn"
            warnings.append("Opex mode missing; default NNN applied.")
    scenario.setdefault("base_opex_psf_yr", 10.0)
    scenario.setdefault("base_year_opex_psf_yr", 10.0)
    scenario.setdefault("opex_growth", 0.03)
    scenario.setdefault("discount_rate_annual", 0.08)
    scenario.setdefault("parking_spaces", 0)
    scenario.setdefault("parking_cost_monthly_per_space", 0.0)
    scenario.setdefault("parking_sales_tax_rate", 0.0825)
    scenario.setdefault("broker_fee", 0.0)
    scenario.setdefault("security_deposit_months", 0.0)
    scenario.setdefault("holdover_months", 0)
    scenario.setdefault("holdover_rent_multiplier", 1.5)
    scenario.setdefault("sublease_income_monthly", 0.0)
    scenario.setdefault("sublease_start_month", 0)
    scenario.setdefault("sublease_duration_months", 0)
    if "one_time_costs" not in scenario:
        scenario["one_time_costs"] = []
    if "termination_option" not in scenario:
        scenario["termination_option"] = None

    for k in list(confidence):
        try:
            confidence[k] = float(confidence[k])
        except (TypeError, ValueError):
            confidence[k] = 0.0

    return scenario, confidence, warnings


def extract_scenario_from_text(text: str, source: str) -> ExtractionResponse:
    """
    Regex prefill + snippet pack + LLM; return ExtractionResponse.
    """
    original_length = len(text)
    if not text or len(text.strip()) < 50:
        scenario = Scenario(
            name="Extracted lease",
            rsf=10000.0,
            commencement=date(2026, 1, 1),
            expiration=date(2031, 1, 31),
            rent_steps=[RentStep(start=0, end=59, rate_psf_yr=0.0)],
            free_rent_months=0,
            ti_allowance_psf=0.0,
            opex_mode=OpexMode.NNN,
            base_opex_psf_yr=10.0,
            base_year_opex_psf_yr=10.0,
            opex_growth=0.03,
            discount_rate_annual=0.08,
        )
        return ExtractionResponse(
            scenario=scenario,
            confidence={},
            warnings=["Document text too short or empty; all values are defaults. Please review."],
            source=source,
            text_length=original_length,
        )

    try:
        prefill = _regex_prefill(text)
        text_for_llm, extra_warnings = _prepare_text_for_llm(text)
        raw: dict | None = None
        llm_error: Exception | None = None
        try:
            raw = _llm_extract_scenario(text_for_llm, prefill=prefill)
        except Exception as e:
            llm_error = e
            logger.warning("[extract] AI extraction failed; using deterministic fallback. err=%s", str(e)[:280])
            print("AI_EXTRACT_FAIL", {"err": str(e)[:400], "type": type(e).__name__}, flush=True)

        if raw is None:
            fallback_scenario, fallback_confidence, fallback_warnings = _heuristic_extract_scenario(
                text,
                prefill=prefill,
                llm_error=llm_error,
            )
            fallback_raw = {
                "scenario": fallback_scenario,
                "confidence": fallback_confidence,
                "warnings": fallback_warnings,
            }
            scenario_dict, confidence, warnings = _apply_safe_defaults(fallback_raw, prefill=prefill)
        else:
            scenario_dict, confidence, warnings = _apply_safe_defaults(raw, prefill=prefill)

        warnings = extra_warnings + warnings
        scenario = Scenario.model_validate(scenario_dict)
        return ExtractionResponse(
            scenario=scenario,
            confidence=confidence,
            warnings=warnings,
            source=source,
            text_length=original_length,
        )
    except Exception as e:
        logger.exception("[extract] unexpected extraction failure; forcing deterministic fallback")
        prefill = _regex_prefill(text) if text else {}
        fallback_scenario, fallback_confidence, fallback_warnings = _heuristic_extract_scenario(
            text,
            prefill=prefill,
            llm_error=e,
        )
        scenario_dict, confidence, warnings = _apply_safe_defaults(
            {"scenario": fallback_scenario, "confidence": fallback_confidence, "warnings": fallback_warnings},
            prefill=prefill,
        )
        scenario = Scenario.model_validate(scenario_dict)
        return ExtractionResponse(
            scenario=scenario,
            confidence=confidence,
            warnings=warnings,
            source=source,
            text_length=original_length,
        )
