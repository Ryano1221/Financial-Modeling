"""
Universal Input Layer — all inputs normalize to CanonicalLease.

Accepts: manual form, pasted text, PDF, Word, Excel, JSON.
Returns: CanonicalLease if confidence >= 0.85; otherwise returns
canonical_lease + confidence_score + missing_fields + clarification_questions.
Frontend must display confirmation step before finalizing when confidence < 0.85.
"""

from __future__ import annotations

from datetime import date, datetime
from enum import Enum
from typing import Any, BinaryIO, Dict, List, Optional, Union

from pydantic import BaseModel, Field

from models import (
    CanonicalLease,
    RentScheduleStep,
    PhaseInStep,
    FreeRentPeriod,
    Scenario,
    LeaseType,
    EscalationType,
    ExpenseStructureType,
)


class InputSource(str, Enum):
    MANUAL = "manual"
    PASTED_TEXT = "pasted_text"
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    JSON = "json"


class NormalizerInput(BaseModel):
    """Input to the normalizer; one of the payloads set."""
    source: InputSource = InputSource.MANUAL
    manual_form: Optional[Dict[str, Any]] = None
    pasted_text: Optional[str] = None
    json_payload: Optional[Dict[str, Any]] = None
    scenario_id: str = ""
    scenario_name: str = ""


class NormalizerResponse(BaseModel):
    """When confidence < 0.85, return this so frontend can show confirmation."""
    canonical_lease: CanonicalLease
    confidence_score: float = Field(ge=0.0, le=1.0)
    missing_fields: List[str] = Field(default_factory=list)
    clarification_questions: List[str] = Field(default_factory=list)


CONFIDENCE_THRESHOLD = 0.85


def _date_from_iso(s: str) -> date:
    """Parse YYYY-MM-DD to date."""
    if isinstance(s, date):
        return s
    return datetime.fromisoformat(s.replace("Z", "+00:00").split("T")[0]).date()


def _scenario_to_canonical(scenario: Scenario, scenario_id: str = "", scenario_name: str = "") -> CanonicalLease:
    """Convert legacy Scenario to CanonicalLease."""
    term = scenario.term_months
    rent_schedule = [
        RentScheduleStep(
            start_month=s.start,
            end_month=s.end,
            rent_psf_annual=s.rate_psf_yr,
            escalation_type=EscalationType.FIXED,
            escalation_value=0.0,
        )
        for s in scenario.rent_steps
    ]
    expense_type = ExpenseStructureType.BASE_YEAR if scenario.opex_mode.value == "base_year" else ExpenseStructureType.NNN
    return CanonicalLease(
        scenario_id=scenario_id or scenario.name,
        scenario_name=scenario_name or scenario.name,
        premises_name=scenario.name,
        address="",
        building_name="",
        suite="",
        floor="",
        rsf=scenario.rsf,
        lease_type=LeaseType.NNN if scenario.opex_mode.value == "nnn" else LeaseType.NNN,
        commencement_date=scenario.commencement,
        expiration_date=scenario.expiration,
        term_months=term,
        free_rent_months=scenario.free_rent_months,
        free_rent_scope="base",
        discount_rate_annual=scenario.discount_rate_annual,
        notes="",
        rent_schedule=rent_schedule,
        phase_in_schedule=[],
        opex_psf_year_1=scenario.base_opex_psf_yr,
        opex_growth_rate=scenario.opex_growth,
        expense_stop_psf=scenario.base_year_opex_psf_yr,
        base_year=None,
        pro_rata_share=1.0,
        expense_structure_type=expense_type,
        parking_ratio=0.0,
        parking_count=scenario.parking_spaces,
        parking_rate_monthly=scenario.parking_cost_monthly_per_space,
        parking_escalation_rate=0.0,
        ti_allowance_psf=scenario.ti_allowance_psf,
        ti_total=scenario.ti_allowance_psf * scenario.rsf,
        landlord_work_value=scenario.ti_allowance_psf * scenario.rsf,
        tenant_capex_total=0.0,
        amortized_ti_flag=False,
        amortization_rate=0.0,
        amortization_term_months=0,
        free_rent_periods=[
            FreeRentPeriod(start_month=0, end_month=max(0, scenario.free_rent_months - 1))
        ] if scenario.free_rent_months > 0 else [],
        rent_abatements=[],
        moving_allowance=0.0,
        other_concessions=[],
        renewal_options=[],
    )


def _dict_to_canonical(data: Dict[str, Any], scenario_id: str = "", scenario_name: str = "") -> CanonicalLease:
    """Build CanonicalLease from a flat or nested dict (e.g. API payload or manual form)."""
    def get(key: str, default: Any = None) -> Any:
        return data.get(key, data.get(key.replace("_", " "), default))

    comm = get("commencement_date") or get("commencement")
    exp = get("expiration_date") or get("expiration")
    if isinstance(comm, str):
        comm = _date_from_iso(comm)
    if isinstance(exp, str):
        exp = _date_from_iso(exp)

    rent_schedule: List[RentScheduleStep] = []
    for step in data.get("rent_schedule", data.get("rent_steps", [])):
        if isinstance(step, dict):
            rent_schedule.append(RentScheduleStep(
                start_month=int(step.get("start_month", step.get("start", 0))),
                end_month=int(step.get("end_month", step.get("end", 0))),
                rent_psf_annual=float(step.get("rent_psf_annual", step.get("rate_psf_yr", 0))),
                escalation_type=EscalationType(step.get("escalation_type", "fixed")),
                escalation_value=float(step.get("escalation_value", 0)),
            ))
        else:
            rent_schedule.append(step)

    if not rent_schedule and (get("rate_psf_yr") or get("rent_psf_annual")) is not None:
        term = int(get("term_months", 60))
        rate = float(get("rate_psf_yr") or get("rent_psf_annual", 0))
        rent_schedule = [RentScheduleStep(start_month=0, end_month=max(0, term - 1), rent_psf_annual=rate)]

    free_rent_months = int(get("free_rent_months", 0) or 0)
    free_rent_start = int(get("free_rent_start_month", 0) or 0)
    free_rent_end = int(get("free_rent_end_month", max(0, free_rent_start + free_rent_months - 1)) or max(0, free_rent_start + free_rent_months - 1))
    free_rent_scope = str(get("free_rent_scope", get("free_rent_abatement_type", "base")) or "base").strip().lower()
    if free_rent_scope not in {"base", "gross"}:
        free_rent_scope = "base"
    free_rent_periods: List[FreeRentPeriod] = []
    raw_free_periods = data.get("free_rent_periods", [])
    for step in raw_free_periods or []:
        if isinstance(step, dict):
            free_rent_periods.append(
                FreeRentPeriod(
                    start_month=int(step.get("start_month", 0)),
                    end_month=int(step.get("end_month", 0)),
                )
            )
        elif isinstance(step, FreeRentPeriod):
            free_rent_periods.append(step)
    if not free_rent_periods and free_rent_months > 0:
        start = max(0, free_rent_start)
        end = max(start, free_rent_end)
        free_rent_periods = [FreeRentPeriod(start_month=start, end_month=end)]

    phase_in_schedule: List[PhaseInStep] = []
    raw_phase_in = data.get("phase_in_schedule", data.get("phase_in_steps", []))
    for step in raw_phase_in or []:
        if isinstance(step, dict):
            phase_in_schedule.append(
                PhaseInStep(
                    start_month=int(step.get("start_month", step.get("start", 0))),
                    end_month=int(step.get("end_month", step.get("end", 0))),
                    rsf=float(step.get("rsf", 0) or 0),
                )
            )
        elif isinstance(step, PhaseInStep):
            phase_in_schedule.append(step)

    lease_type_str = (get("lease_type") or get("opex_mode") or "NNN").strip()
    try:
        lt = LeaseType(lease_type_str)
    except ValueError:
        lt = LeaseType.NNN
    exp_type = (get("expense_structure_type") or "nnn").strip().lower()
    try:
        est = ExpenseStructureType(exp_type)
    except ValueError:
        est = ExpenseStructureType.NNN

    building_name = str(get("building_name", "") or "").strip()
    suite = str(get("suite", "") or "").strip()
    floor = str(get("floor", "") or "").strip()
    premises = str(get("premises_name", get("name", "") or "")).strip()
    if building_name and suite and not premises:
        premises = f"{building_name} Suite {suite}"
    elif building_name and floor and not premises:
        premises = f"{building_name} Floor {floor}"

    return CanonicalLease(
        scenario_id=scenario_id or get("scenario_id", ""),
        scenario_name=scenario_name or get("scenario_name", get("name", "Unnamed")),
        premises_name=premises,
        address=get("address", ""),
        building_name=building_name,
        suite=suite,
        floor=floor,
        rsf=float(get("rsf", 0) or 0),
        lease_type=lt,
        commencement_date=comm or date(2026, 1, 1),
        expiration_date=exp or date(2031, 1, 1),
        term_months=int(get("term_months", 60)),
        free_rent_months=free_rent_months,
        free_rent_scope=free_rent_scope,  # type: ignore[arg-type]
        discount_rate_annual=float(get("discount_rate_annual", 0.08) or 0.08),
        notes=str(get("notes", "")),
        rent_schedule=rent_schedule,
        phase_in_schedule=phase_in_schedule,
        opex_psf_year_1=float(get("opex_psf_year_1", get("base_opex_psf_yr", 0)) or 0),
        opex_growth_rate=float(get("opex_growth_rate", get("opex_growth", 0)) or 0),
        expense_stop_psf=float(get("expense_stop_psf", get("base_year_opex_psf_yr", 0)) or 0),
        base_year=get("base_year"),
        pro_rata_share=float(get("pro_rata_share", 1) or 1),
        expense_structure_type=est,
        parking_ratio=float(get("parking_ratio", 0) or 0),
        parking_count=int(get("parking_count", get("parking_spaces", 0)) or 0),
        parking_rate_monthly=float(get("parking_rate_monthly", get("parking_cost_monthly_per_space", 0)) or 0),
        parking_escalation_rate=float(get("parking_escalation_rate", 0) or 0),
        ti_allowance_psf=float(get("ti_allowance_psf", 0) or 0),
        ti_total=float(get("ti_total", 0) or 0),
        landlord_work_value=float(get("landlord_work_value", 0) or 0),
        tenant_capex_total=float(get("tenant_capex_total", 0) or 0),
        amortized_ti_flag=bool(get("amortized_ti_flag", False)),
        amortization_rate=float(get("amortization_rate", 0) or 0),
        amortization_term_months=int(get("amortization_term_months", 0) or 0),
        free_rent_periods=free_rent_periods,
        rent_abatements=[],
        moving_allowance=float(get("moving_allowance", 0) or 0),
        other_concessions=list(get("other_concessions", []) or []),
        renewal_options=[],
    )


def _compute_confidence_and_missing(lease: CanonicalLease) -> tuple[float, List[str], List[str]]:
    """Compute aggregate confidence, missing fields, and clarification questions."""
    missing: List[str] = []
    questions: List[str] = []
    has_premises = (
        bool((lease.premises_name or "").strip())
        or bool((lease.building_name or "").strip())
        or bool((lease.suite or "").strip())
        or bool((lease.floor or "").strip())
    )
    if not has_premises:
        missing.append("premises_name")
        questions.append("What is the building name or suite?")
    if lease.rsf <= 0:
        missing.append("rsf")
        questions.append("What is the rentable square footage?")
    if not lease.rent_schedule:
        missing.append("rent_schedule")
        questions.append("Please provide the base rent schedule (rate and period).")
    if lease.term_months <= 0:
        missing.append("term_months")
        questions.append("What is the lease term in months?")
    n_missing = len(missing)
    confidence = max(0.0, 1.0 - (n_missing * 0.15))
    return confidence, missing, questions


def normalize_input(
    source: InputSource,
    payload: Union[Dict[str, Any], Scenario, str, BinaryIO, bytes],
    scenario_id: str = "",
    scenario_name: str = "",
) -> Union[CanonicalLease, NormalizerResponse]:
    """
    Normalize any input to CanonicalLease.
    If confidence < 0.85, returns NormalizerResponse for frontend confirmation.
    """
    canonical: Optional[CanonicalLease] = None
    confidence = 1.0

    if source == InputSource.MANUAL or source == InputSource.JSON:
        if isinstance(payload, Scenario):
            canonical = _scenario_to_canonical(payload, scenario_id, scenario_name)
        elif isinstance(payload, dict):
            canonical = _dict_to_canonical(payload, scenario_id, scenario_name)
        else:
            data = getattr(payload, "model_dump", lambda: payload)() if hasattr(payload, "model_dump") else {}
            canonical = _dict_to_canonical(data, scenario_id, scenario_name)
        confidence, missing, questions = _compute_confidence_and_missing(canonical)
    elif source == InputSource.PASTED_TEXT:
        try:
            from scenario_extract import extract_scenario_from_text
            extraction = extract_scenario_from_text(str(payload), "pasted_text")
            if extraction and extraction.scenario:
                canonical = _scenario_to_canonical(extraction.scenario, scenario_id, scenario_name)
                confidence = 0.78
            else:
                canonical = _dict_to_canonical({}, scenario_id, scenario_name)
                confidence = 0.5
        except Exception:
            canonical = _dict_to_canonical({}, scenario_id, scenario_name)
            confidence = 0.5
        confidence, missing, questions = _compute_confidence_and_missing(canonical)
    elif source in (InputSource.PDF, InputSource.WORD, InputSource.EXCEL):
        try:
            text = ""
            if source == InputSource.PDF:
                from lease_extract import extract_text_from_pdf
                text = extract_text_from_pdf(payload) if hasattr(payload, "read") else extract_text_from_pdf(open(payload, "rb"))  # type: ignore
            elif source == InputSource.WORD:
                try:
                    import docx  # type: ignore
                    doc = docx.Document(payload) if hasattr(payload, "read") else docx.Document(open(payload, "rb"))  # type: ignore
                    text = "\n".join(p.text for p in doc.paragraphs)
                except Exception:
                    pass
            if text and len(text.strip()) > 50:
                from scenario_extract import extract_scenario_from_text
                extraction = extract_scenario_from_text(text, "pdf_text" if source == InputSource.PDF else "docx")
                if extraction and extraction.scenario:
                    canonical = _scenario_to_canonical(extraction.scenario, scenario_id, scenario_name)
                    confidence = 0.78
                else:
                    canonical = _dict_to_canonical({}, scenario_id, scenario_name)
                    confidence = 0.5
            else:
                canonical = _dict_to_canonical({}, scenario_id, scenario_name)
                confidence = 0.4
        except Exception:
            canonical = _dict_to_canonical({}, scenario_id, scenario_name)
            confidence = 0.4
        confidence, missing, questions = _compute_confidence_and_missing(canonical)
    else:
        canonical = _dict_to_canonical(payload if isinstance(payload, dict) else {}, scenario_id, scenario_name)
        confidence, missing, questions = _compute_confidence_and_missing(canonical)

    if confidence >= CONFIDENCE_THRESHOLD:
        return canonical
    return NormalizerResponse(
        canonical_lease=canonical,
        confidence_score=confidence,
        missing_fields=missing,
        clarification_questions=questions,
    )
